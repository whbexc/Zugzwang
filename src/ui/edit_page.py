from __future__ import annotations
_SCROLLBAR_STYLE = """
    QScrollBar:vertical {
        background: #1a1a1a;
        width: 4px;
        margin: 0;
    }
    QScrollBar::handle:vertical {
        background: #333333;
        border-radius: 2px;
        min-height: 20px;
    }
    QScrollBar::handle:vertical:hover {
        background: #444444;
    }
    QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {
        height: 0px;
        background: transparent;
    }
    QScrollBar::add-page:vertical, QScrollBar::sub-page:vertical {
        background: transparent;
    }
    QScrollBar:horizontal {
        height: 0px;
        background: transparent;
    }
"""
"""
ZUGZWANG - Motivation Letter Edit Workspace
Auto-generates an Anschreiben from persisted lead data.

Improvements over v1:
  - Letter word / char counter in editor header
  - Per-lead letter history (undo to previous saved version)
  - Batch export: export letters for all filtered leads at once
  - Export as .docx (python-docx) in addition to .txt
  - Find & Replace toolbar inside the editor (Ctrl+H)
  - Placeholder audit: highlights any un-filled {{...}} remaining
  - Letter preview mode: read-only styled view toggled by a button
  - Auto-save on lead switch (debounced, 800ms)
  - Template variable list shown inline in the Edit Template dialog
  - Sender block injected automatically from settings (if set)
  - Status bar shows word count + last saved timestamp
"""


import re
import sqlite3
from dataclasses import dataclass, field
from datetime import date, datetime
from pathlib import Path
from typing import Optional

from PySide6.QtCore import Qt, QTimer, QSize, Signal, QCoreApplication, QStandardPaths, QUrl, QThread, QEvent, QPoint
from PySide6.QtGui import (
    QDesktopServices,
    QGuiApplication,
    QTextCharFormat,
    QColor,
    QFont,
    QTextCursor,
    QTextBlockFormat,
    QKeySequence,
    QShortcut,
    QPalette,
)
from PySide6.QtWidgets import (
    QDialog,
    QFileDialog,
    QFrame,
    QHBoxLayout,
    QLabel,
    QListWidget,
    QListWidgetItem,
    QSizePolicy,
    QTextEdit,
    QVBoxLayout,
    QWidget,
    QScrollArea,
    QAbstractItemView,
    QLineEdit,
    QCheckBox,
    QMessageBox,
    QPushButton,
    QSplitter,
    QSpacerItem,
    QGraphicsOpacityEffect,
)
from qfluentwidgets import InfoBar, LineEdit, PushButton, InfoBarPosition, CaptionLabel, FluentIcon, IconWidget, RoundMenu, Action


class SegmentTabButton(QPushButton):
    def __init__(self, label: str, parent=None):
        super().__init__(parent)
        self._base_label = label
        self._count = 0
        self._is_active = False
        self.setFixedHeight(28)
        self.setCheckable(True)
        self.setCursor(Qt.PointingHandCursor)

        layout = QHBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)
        layout.setAlignment(Qt.AlignCenter)

        self._display = QLabel(label)
        self._display.setAttribute(Qt.WA_TransparentForMouseEvents)
        self._display.setAlignment(Qt.AlignCenter)
        self._display.setTextFormat(Qt.RichText)
        self._display.setStyleSheet(
            "background: transparent; border: none;"
        )
        layout.addWidget(self._display)

        self.active_style = """
            QPushButton {
                background: transparent;
                border: none;
                border-bottom: 2px solid #4CAF50;
                border-radius: 0px;
            }
        """
        self.inactive_style = """
            QPushButton {
                background: transparent;
                border: none;
                border-bottom: 2px solid transparent;
                border-radius: 0px;
            }
            QPushButton:hover {
                background: rgba(255, 255, 255, 0.03);
            }
        """
        self.setStyleSheet(self.inactive_style)
        self._refresh_label()

    def set_active(self, active: bool):
        self._is_active = active
        self.setChecked(active)
        self.setStyleSheet(self.active_style if active else self.inactive_style)
        self._refresh_label()

    def set_count(self, count: int):
        self._count = count
        self._refresh_label()

    def _refresh_label(self):
        if self._is_active:
            label_color = "#ffffff"
            count_color = "#888888"
            weight = 500
        else:
            label_color = "#aaaaaa"
            count_color = "#555555"
            weight = 400
        count_str = f" {self._count}" if self._count > 0 else ""
        self._display.setText(
            f"<span style='font-family: system-ui, -apple-system, sans-serif; "
            f"font-size: 12px; font-weight: {weight}; color: {label_color};'>"
            f"{self._base_label}</span>"
            f"<span style='font-family: system-ui, -apple-system, sans-serif; "
            f"font-size: 11px; font-weight: {weight}; color: {count_color};'>"
            f"{count_str}</span>"
        )

from ..core.config import get_exports_dir, get_memory_db_path
from ..core.models import LeadRecord
from ..services.export_service import ExportService
from ..utils.db_worker import run_in_thread
from ..core.events import EventBus, event_bus
from .event_bridge import event_bridge
from .theme import Theme
from .icons import _render_tinted_icon


class PdfDetectWorker(QThread):
    detection_done = Signal(int, int, bool, str)

    def __init__(self, path: Path, parent=None):
        super().__init__(parent)
        self.path = path

    def run(self):
        try:
            import pypdf
        except ImportError:
            self.detection_done.emit(0, -1, False, "Install pypdf:  pip install pypdf")
            return

        try:
            reader = pypdf.PdfReader(self.path)
            if reader.is_encrypted:
                self.detection_done.emit(0, -1, False, "PDF is password-protected. Please unlock it first.")
                return
            
            num_pages = len(reader.pages)
            detected_idx = 1 if num_pages > 1 else 0
            was_detected = False

            # Try to detect Anschreiben
            for i, page in enumerate(reader.pages):
                try:
                    text = page.extract_text() or ""
                    if any(kw in text for kw in ["Anschreiben", "Motivation", "Sehr geehrte", "Mit freundlichen Grüßen"]):
                        detected_idx = i
                        was_detected = True
                        break
                except Exception:
                    pass

            self.detection_done.emit(num_pages, detected_idx, was_detected, "")
        except Exception as e:
            self.detection_done.emit(0, -1, False, str(e))


GERMAN_MONTHS = [
    "Januar", "Februar", "März", "April", "Mai", "Juni",
    "Juli", "August", "September", "Oktober", "November", "Dezember",
]

PLACEHOLDER_REFERENCE = [
    ("{{ANREDE}}",  "Sehr geehrte Frau Müller, / Sehr geehrter Herr …"),
    ("{{FIRMA}}",   "Unternehmensname"),
    ("{{ORT}}",     "Stadt"),
    ("{{PLZ}}",     "Postleitzahl"),
    ("{{BERUF}}",   "Ausbildungsberuf / Stelle"),
    ("{{DATUM}}",   "04. Juni 2026"),
    ("{{SENDER_NAME}}",    "Ihr Name (aus Einstellungen)"),
    ("{{SENDER_ADDRESS}}", "Ihre Straße + Hausnummer"),
    ("{{SENDER_CITY}}",    "Ihre PLZ + Stadt"),
    ("{{SENDER_PHONE}}",   "Ihre Telefonnummer"),
    ("{{SENDER_EMAIL}}",   "Ihre E-Mail-Adresse"),
]


# ─────────────────────────────────────────────────────────────────────────────
# Data model
# ─────────────────────────────────────────────────────────────────────────────

@dataclass
class LetterState:
    generated_at: str = ""
    edited_at: str = ""
    sent_at: str = ""
    letter_text: str = ""
    previous_text: str = ""          # one-level undo (last saved version)
    last_saved_ts: str = ""          # human-readable "saved HH:MM"
    is_discarded: int = 0            # hidden from edit page but kept in db


# ─────────────────────────────────────────────────────────────────────────────
# Lead list row widget
# ─────────────────────────────────────────────────────────────────────────────

class ElidedLabel(QLabel):
    """A QLabel that automatically elides text when it exceeds the width of the label."""
    def __init__(self, text="", parent=None):
        super().__init__("", parent)
        self._full_text = text
        self.setText(text)

    def setText(self, text):
        self._full_text = text
        self.update_elided()

    def resizeEvent(self, event):
        super().resizeEvent(event)
        self.update_elided()

    def update_elided(self):
        fm = self.fontMetrics()
        elided = fm.elidedText(self._full_text, Qt.ElideRight, max(0, self.width() - 4))
        if elided != super().text():
            super().setText(elided)


class LeadRowWidget(QFrame):
    """Compact lead browser row with status badge."""

    def __init__(self, record: LeadRecord, state: LetterState | None = None):
        super().__init__()
        self.record_id = record.id
        self.setObjectName("leadRow")
        self.setProperty("selected", False)
        self.setFixedHeight(70)
        self.setCursor(Qt.PointingHandCursor)

        root = QVBoxLayout(self)
        root.setContentsMargins(16, 6, 12, 6)
        root.setSpacing(2)

        top = QHBoxLayout()
        top.setContentsMargins(0, 0, 0, 0)
        top.setSpacing(5)

        company = ElidedLabel(record.company_name or "Unknown company")
        company.setStyleSheet("""
            color: #FFFFFF;
            font-family: 'PT Root UI', sans-serif;
            font-size: 13px;
            font-weight: 600;
            background: transparent;
            border: none;
            padding: 0px;
        """)
        top.addWidget(company, 1)

        # Show GEN badge only when generated (not sent)
        if state and state.generated_at and not state.sent_at:
            badge = QLabel("GEN")
            badge.setAlignment(Qt.AlignCenter)
            badge.setFixedHeight(18)
            badge.setSizePolicy(QSizePolicy.Fixed, QSizePolicy.Fixed)
            badge.setStyleSheet("""
                color: #a0a0a5;
                background: rgba(255, 255, 255, 0.05);
                border: 1px solid rgba(255, 255, 255, 0.1);
                border-radius: 4px;
                font-family: 'SF Mono', monospace;
                font-size: 9px;
                font-weight: 600;
                padding: 2px 6px;
            """)
            top.addWidget(badge, 0, Qt.AlignVCenter)

        root.addLayout(top)

        job    = record.job_title or record.category or "—"
        city   = record.city or ""
        sub_text = f"{job} · {city}" if city else job
        meta   = ElidedLabel(sub_text)
        meta.setStyleSheet("""
            color: #6e6e73;
            font-family: 'PT Root UI', sans-serif;
            font-size: 11px;
            font-weight: 400;
            background: transparent;
            border: none;
            padding: 0px;
        """)
        root.addWidget(meta)

        self._apply_style()

    def set_selected(self, selected: bool):
        if self.property("selected") == selected:
            return
        self.setProperty("selected", selected)
        self._apply_style()

    def _apply_style(self):
        if self.property("selected"):
            self.setStyleSheet("""
                QFrame#leadRow {
                    background: rgba(10, 132, 255, 0.10);
                    border: none;
                    border-left: 2px solid #0A84FF;
                    border-radius: 0px;
                }
            """)
        else:
            self.setStyleSheet("""
                QFrame#leadRow {
                    background: transparent;
                    border: none;
                    border-left: 2px solid transparent;
                    border-radius: 0px;
                }
                QFrame#leadRow:hover {
                    background: #2C2C2E;
                }
            """)


# ─────────────────────────────────────────────────────────────────────────────
# Find & Replace toolbar
# ─────────────────────────────────────────────────────────────────────────────

class FindReplaceBar(QFrame):
    """Compact inline find/replace toolbar."""

    def __init__(self, editor: QTextEdit, parent=None):
        super().__init__(parent)
        self._editor = editor
        self.setObjectName("findBar")
        self.setFixedHeight(44)
        self.setStyleSheet("""
            QFrame#findBar {
                background: #1a1a1a;
                border: 1px solid #2a2a2a;
                border-radius: 4px;
            }
        """)

        layout = QHBoxLayout(self)
        layout.setContentsMargins(10, 0, 10, 0)
        layout.setSpacing(8)

        input_ss = (
            "QLineEdit { background: #111111; border: 1px solid #2a2a2a; "
            "border-radius: 4px; color: #cccccc; padding: 0 8px; font-size: 12px; }"
            "QLineEdit:focus { border-color: #4CAF50; }"
        )

        self._find_input = QLineEdit()
        self._find_input.setPlaceholderText("Find…")
        self._find_input.setFixedHeight(28)
        self._find_input.setStyleSheet(input_ss)
        self._find_input.returnPressed.connect(self._find_next)

        self._replace_input = QLineEdit()
        self._replace_input.setPlaceholderText("Replace…")
        self._replace_input.setFixedHeight(28)
        self._replace_input.setStyleSheet(input_ss)

        self._case_cb = QCheckBox("Aa")
        self._case_cb.setStyleSheet(
            "QCheckBox { color: #AEAEB2; font-size: 11px; }"
            "QCheckBox::indicator { width: 14px; height: 14px; }"
        )

        btn_ss = (
            "QPushButton { background: #1a1a1a; border: 1px solid #2a2a2a; border-radius: 4px; "
            "color: #999999; font-size: 11px; padding: 3px 10px; }"
            "QPushButton:hover { background: #222222; border-color: #444444; color: #eeeeee; }"
            "QPushButton:pressed { background: #4CAF50; color: #ffffff; }"
        )

        btn_next    = self._small_btn("Next",       btn_ss, self._find_next)
        btn_prev    = self._small_btn("Prev",       btn_ss, self._find_prev)
        btn_replace = self._small_btn("Replace",    btn_ss, self._replace_one)
        btn_all     = self._small_btn("Replace All",btn_ss, self._replace_all)
        btn_close   = self._small_btn("✕",          btn_ss, self.hide)

        self._match_label = QLabel("")
        self._match_label.setStyleSheet(
            "color: #636366; font-size: 10px; background: transparent;"
        )

        layout.addWidget(QLabel("🔍", styleSheet="color:#8E8E93; background:transparent;"))
        layout.addWidget(self._find_input, 2)
        layout.addWidget(self._replace_input, 2)
        layout.addWidget(self._case_cb)
        layout.addWidget(btn_next)
        layout.addWidget(btn_prev)
        layout.addWidget(btn_replace)
        layout.addWidget(btn_all)
        layout.addWidget(self._match_label, 1)
        layout.addWidget(btn_close)

    def _small_btn(self, text, ss, slot):
        b = PushButton(text)
        b.setFixedHeight(26)
        b.setStyleSheet(ss)
        b.clicked.connect(slot)
        return b

    def show_and_focus(self):
        self.show()
        self._find_input.setFocus()
        self._find_input.selectAll()

    # ── find/replace logic ──────────────────────────────────────────────────

    def _flags(self):
        f = QTextDocument.FindFlags() if False else re.IGNORECASE
        return 0 if self._case_cb.isChecked() else re.IGNORECASE

    def _find_next(self, backward=False):
        needle = self._find_input.text()
        if not needle:
            return
        flags = QTextEdit.ExtraSelection
        doc = self._editor.document()
        cursor = self._editor.textCursor()
        find_flags = (
            self._editor.document().defaultTextOption().flags()
        )
        # use Qt built-in find
        from PySide6.QtGui import QTextDocument
        qt_flags = QTextDocument.FindFlag(0)
        if self._case_cb.isChecked():
            qt_flags |= QTextDocument.FindCaseSensitively
        if backward:
            qt_flags |= QTextDocument.FindBackward

        found = self._editor.find(needle, qt_flags)
        if not found:
            # wrap around
            cursor = self._editor.textCursor()
            cursor.movePosition(
                QTextCursor.End if backward else QTextCursor.Start
            )
            self._editor.setTextCursor(cursor)
            found = self._editor.find(needle, qt_flags)

        # count matches
        count = len(re.findall(
            re.escape(needle),
            self._editor.toPlainText(),
            0 if self._case_cb.isChecked() else re.IGNORECASE,
        ))
        self._match_label.setText(f"{count} match{'es' if count != 1 else ''}")

    def _find_prev(self):
        self._find_next(backward=True)

    def _replace_one(self):
        needle      = self._find_input.text()
        replacement = self._replace_input.text()
        if not needle:
            return
        cursor = self._editor.textCursor()
        if cursor.hasSelection() and cursor.selectedText() == needle:
            cursor.insertText(replacement)
        self._find_next()

    def _replace_all(self):
        needle      = self._find_input.text()
        replacement = self._replace_input.text()
        if not needle:
            return
        text  = self._editor.toPlainText()
        flags = 0 if self._case_cb.isChecked() else re.IGNORECASE
        new_text, n = re.subn(re.escape(needle), replacement, text, flags=flags)
        if n:
            self._editor.setPlainText(new_text)
            self._match_label.setText(f"Replaced {n}×")


# ─────────────────────────────────────────────────────────────────────────────
# Placeholder highlighter
# ─────────────────────────────────────────────────────────────────────────────

def _highlight_placeholders(editor: QTextEdit):
    """
    Highlight any remaining {{...}} in red so the user knows they are un-filled.
    """
    fmt = QTextCharFormat()
    fmt.setBackground(QColor("#5C1A1A"))
    fmt.setForeground(QColor("#FF6B6B"))

    extras: list[QTextEdit.ExtraSelection] = []
    text  = editor.toPlainText()
    for m in re.finditer(r"\{\{[A-Z_]+\}\}", text):
        sel           = QTextEdit.ExtraSelection()
        sel.format    = fmt
        cursor        = editor.textCursor()
        cursor.setPosition(m.start())
        cursor.setPosition(m.end(), QTextCursor.KeepAnchor)
        sel.cursor    = cursor
        extras.append(sel)

    editor.setExtraSelections(extras)
    return len(extras)


# ─────────────────────────────────────────────────────────────────────────────
# PDF Drop Zone Widget
# ─────────────────────────────────────────────────────────────────────────────

class PDFDropZone(QFrame):
    file_dropped = Signal(str)
    clicked = Signal()

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setAcceptDrops(True)
        self.setCursor(Qt.PointingHandCursor)
        self.setObjectName("pdfDropZone")
        
        self.setStyleSheet("""
            QFrame#pdfDropZone {
                background: rgba(255,255,255,0.035);
                border: 0.5px dashed rgba(255,255,255,0.18);
                border-radius: 9px;
            }
            QFrame#pdfDropZone:hover {
                border-color: #0A84FF;
            }
        """)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(10, 10, 10, 10)
        layout.setAlignment(Qt.AlignCenter)
        self.label = QLabel("Drop Bewerbung PDF here\nor click to browse")
        self.label.setAlignment(Qt.AlignCenter)
        self.label.setStyleSheet(
            "color: #8E8E93; font-family: system-ui, -apple-system, sans-serif; "
            "font-size: 11px; font-weight: 400; background: transparent; border: none;"
        )
        layout.addWidget(self.label)

    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            urls = event.mimeData().urls()
            if len(urls) == 1 and urls[0].toLocalFile().lower().endswith(".pdf"):
                event.acceptProposedAction()
                return
        event.ignore()

    def dragMoveEvent(self, event):
        event.acceptProposedAction()

    def dropEvent(self, event):
        urls = event.mimeData().urls()
        if urls:
            path = urls[0].toLocalFile()
            self.file_dropped.emit(path)
            event.acceptProposedAction()

    def mousePressEvent(self, event):
        if event.button() == Qt.LeftButton:
            self.clicked.emit()
            event.accept()
        else:
            super().mousePressEvent(event)


# ─────────────────────────────────────────────────────────────────────────────
# Main page
# ─────────────────────────────────────────────────────────────────────────────

class MappeButton(QPushButton):
    def __init__(self, icon_name: str, tooltip: str, parent=None):
        super().__init__(parent)
        self.icon_name = icon_name
        self.setFixedSize(32, 32)
        self.setCursor(Qt.PointingHandCursor)
        self.setToolTip(tooltip)
        
        self.normal_icon = _render_tinted_icon(icon_name, 14, "#888888")
        self.hover_icon = _render_tinted_icon(icon_name, 14, "#cccccc")
        self.disabled_icon = _render_tinted_icon(icon_name, 14, "#333333")
        
        self.setIcon(self.normal_icon)
        self.setIconSize(QSize(14, 14))
        
        self.setStyleSheet("""
            QPushButton {
                background: #1a1a1a;
                border: 1px solid #2a2a2a;
                border-radius: 4px;
            }
            QPushButton:hover {
                border-color: #444444;
                background: #1f1f1f;
            }
            QPushButton:pressed {
                background: #252525;
            }
            QPushButton:disabled {
                background: #111111;
                border-color: #1a1a1a;
            }
        """)

    def enterEvent(self, event):
        if self.isEnabled():
            self.setIcon(self.hover_icon)
        super().enterEvent(event)

    def leaveEvent(self, event):
        if self.isEnabled():
            self.setIcon(self.normal_icon)
        super().leaveEvent(event)

    def changeEvent(self, event):
        if event.type() == QEvent.EnabledChange:
            if not self.isEnabled():
                self.setIcon(self.disabled_icon)
            else:
                self.setIcon(self.normal_icon)
        super().changeEvent(event)


from PySide6.QtWidgets import QSplitterHandle, QDialog, QVBoxLayout, QHBoxLayout, QLabel, QPushButton, QWidget
from PySide6.QtCore import Qt
from PySide6.QtGui import QPainter, QColor

class AppleConfirmDialog(QDialog):
    def __init__(self, title: str, text: str, confirm_text: str = "Delete", confirm_color: str = "#FF453A", parent=None):
        super().__init__(parent)
        self.setWindowFlags(Qt.Dialog | Qt.FramelessWindowHint | Qt.WindowSystemMenuHint)
        self.setAttribute(Qt.WA_TranslucentBackground)
        self.setFixedSize(320, 160)
        
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        bg = QWidget()
        bg.setStyleSheet("QWidget { background: rgba(40, 40, 40, 0.95); border: 1px solid rgba(255, 255, 255, 0.1); border-radius: 14px; }")
        bg_layout = QVBoxLayout(bg)
        bg_layout.setContentsMargins(20, 20, 20, 20)
        bg_layout.setSpacing(8)

        title_lbl = QLabel(title)
        title_lbl.setStyleSheet("color: #FFFFFF; font-family: 'SF Pro Text', 'PT Root UI', sans-serif; font-size: 16px; font-weight: 600; background: transparent; border: none;")
        title_lbl.setAlignment(Qt.AlignCenter)
        bg_layout.addWidget(title_lbl)

        text_lbl = QLabel(text)
        text_lbl.setWordWrap(True)
        text_lbl.setAlignment(Qt.AlignCenter)
        text_lbl.setStyleSheet("color: rgba(255, 255, 255, 0.7); font-family: 'SF Pro Text', 'PT Root UI', sans-serif; font-size: 13px; font-weight: 400; background: transparent; border: none;")
        bg_layout.addWidget(text_lbl, 1)

        btn_layout = QHBoxLayout()
        btn_layout.setContentsMargins(0, 8, 0, 0)
        btn_layout.setSpacing(10)

        btn_cancel = QPushButton("Cancel")
        btn_cancel.setCursor(Qt.PointingHandCursor)
        btn_cancel.setFixedHeight(32)
        btn_cancel.setStyleSheet("QPushButton { background: rgba(255, 255, 255, 0.1); color: white; border: none; border-radius: 6px; font-family: 'SF Pro Text', 'PT Root UI', sans-serif; font-size: 13px; font-weight: 500; } QPushButton:hover { background: rgba(255, 255, 255, 0.15); }")
        btn_cancel.clicked.connect(self.reject)
        
        btn_delete = QPushButton(confirm_text)
        btn_delete.setCursor(Qt.PointingHandCursor)
        btn_delete.setFixedHeight(32)
        btn_delete.setStyleSheet(f"QPushButton {{ background: {confirm_color}; color: white; border: none; border-radius: 6px; font-family: 'SF Pro Text', 'PT Root UI', sans-serif; font-size: 13px; font-weight: 600; }} QPushButton:hover {{ background: {confirm_color}CC; }}")
        btn_delete.clicked.connect(self.accept)

        btn_layout.addWidget(btn_cancel)
        btn_layout.addWidget(btn_delete)
        
        bg_layout.addLayout(btn_layout)
        layout.addWidget(bg)

class CustomSplitterHandle(QSplitterHandle):
    def paintEvent(self, event):
        painter = QPainter(self)
        painter.fillRect(self.rect(), QColor("#3A3A3C"))

class EditPage(QWidget):
    """Three-panel workspace for auto-generated motivation letters."""
    
    send_emails_to_sender = Signal(list)

    def __init__(self):
        super().__init__()
        self.setObjectName("editPage")
        self.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)

        self._records:               list[LeadRecord]      = []
        self._states:                dict[str, LetterState] = {}
        self._selected_record:       LeadRecord | None      = None
        self._selected_lead_id:      str | None             = None
        self._selected_row_widget:   LeadRowWidget | None   = None
        self._lead_population_token: int                    = 0
        self._pending_lead_records:  list[LeadRecord]       = []
        self._pending_lead_index:    int                    = 0
        self._rendered_text:         str                    = ""
        self._loading_text:          bool                   = False
        self._preview_mode:          bool                   = False
        self._bewerbung_pdf_path:         Path | None = None
        self._bewerbung_anschreiben_page: int         = 1   # 0-based, default page 2

        self._cached_sender_settings: dict = {}
        self._fetch_sender_settings()
        event_bus.subscribe(EventBus.SETTINGS_CHANGED, self._fetch_sender_settings)
        event_bridge.job_result.connect(self._on_live_result_added)
        event_bridge.db_updated.connect(self._on_db_updated)

        self._signature_image_path: str = ""
        run_in_thread(
            self._do_fetch_signature_path,
            on_result=self._on_signature_path_loaded
        )

        run_in_thread(
            self._do_fetch_persisted_pdf_path,
            on_result=self._on_persisted_pdf_path_loaded
        )

        from ..core.config import get_app_data_dir
        import shutil

        self._template_path = get_app_data_dir() / "templates" / "anschreiben_base.txt"
        if not self._template_path.exists():
            default_path = Path(__file__).resolve().parents[2] / "templates" / "anschreiben_base.txt"
            if default_path.exists():
                self._template_path.parent.mkdir(parents=True, exist_ok=True)
                try:
                    shutil.copy2(default_path, self._template_path)
                except Exception as e:
                    import logging
                    logging.getLogger(__name__).warning(f"Could not copy default template: {e}")
        self._template_text = self._load_template()

        # auto-save debounce timer
        self._autosave_timer = QTimer(self)
        self._autosave_timer.setSingleShot(True)
        self._autosave_timer.setInterval(800)
        self._autosave_timer.timeout.connect(self._autosave)

        # word-count refresh timer
        self._wc_timer = QTimer(self)
        self._wc_timer.setSingleShot(True)
        self._wc_timer.setInterval(300)
        self._wc_timer.timeout.connect(self._update_word_count)

        self._build_ui()
        self._setup_shortcuts()
        QTimer.singleShot(0, self.refresh)

    # ─────────────────────────────────────────────────────────────────────────
    # UI construction
    # ─────────────────────────────────────────────────────────────────────────

    def _build_ui(self):
        self.setStyleSheet("QWidget#editPage { background: transparent; }")
        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)
        root.setSpacing(0)

        # ── page header / top bar ─────────────────────────────────────────────
    def _build_ui(self):
        self.setStyleSheet("QWidget#editPage { background: transparent; }")
        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)
        root.setSpacing(0)

        # Child 1: TopBar
        top_bar = QWidget()
        top_bar.setFixedHeight(64)
        top_bar.setStyleSheet("background: #1C1C1E; border-bottom: 0.5px solid #3A3A3C;")
        
        top_layout = QHBoxLayout(top_bar)
        top_layout.setContentsMargins(16, 0, 12, 0)
        top_layout.setSpacing(0)
        top_layout.setAlignment(Qt.AlignVCenter)

        # LEFT
        left_box = QVBoxLayout()
        left_box.setSpacing(1)
        left_box.setAlignment(Qt.AlignVCenter)

        title_lbl = QLabel("Edit")
        title_lbl.setStyleSheet("""
            color: #FFFFFF;
            font-family: 'PT Root UI', sans-serif;
            font-size: 28px;
            font-weight: 600;
            background: transparent;
            border: none;
        """)

        left_box.addWidget(title_lbl)
        top_layout.addLayout(left_box)

        # SPACER
        top_layout.addSpacerItem(QSpacerItem(0, 0, QSizePolicy.Expanding, QSizePolicy.Minimum))



        root.addWidget(top_bar)

        # Child 2: BodySplitter
        class StyledSplitter(QSplitter):
            def createHandle(self):
                return CustomSplitterHandle(self.orientation(), self)

        workspace = StyledSplitter(Qt.Horizontal)
        workspace.setChildrenCollapsible(False)
        workspace.setHandleWidth(1)

        left_panel = self._build_left_panel()
        left_panel.setFixedWidth(280)
        
        center_panel = self._build_center_panel()
        
        right_panel = self._build_right_panel()
        right_panel.setFixedWidth(280)

        workspace.addWidget(left_panel)
        workspace.addWidget(center_panel)
        workspace.addWidget(right_panel)

        workspace.setStretchFactor(0, 0)
        workspace.setStretchFactor(1, 1)
        workspace.setStretchFactor(2, 0)

        root.addWidget(workspace, 1)

        # Setup standard UI state
        self._set_actions_enabled(False)

    def _build_left_panel(self) -> QWidget:
        left = QWidget()
        left.setStyleSheet("background: #1C1C1E;")
        layout = QVBoxLayout(left)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        # A. Header Row (Removed)

        # B. Search Bar
        search_w = QWidget()
        search_lyt = QVBoxLayout(search_w)
        search_lyt.setContentsMargins(10, 10, 10, 6)
        
        search_container = QWidget()
        search_container.setFixedHeight(32)
        search_container.setStyleSheet("background-color: #2C2C2E; border-radius: 7px;")
        sc_layout = QHBoxLayout(search_container)
        sc_layout.setContentsMargins(9, 0, 9, 0)
        sc_layout.setSpacing(6)
        
        s_icon = QLabel()
        s_icon.setPixmap(_render_tinted_icon("search.svg", 13, "#6E6E73").pixmap(13, 13))
        
        self._search = QLineEdit()
        self._search.setPlaceholderText("Search company, HR, city…")
        self._search.setStyleSheet("""
            QLineEdit {
                background: transparent;
                border: none;
                color: #E5E5EA;
                font-family: system-ui, -apple-system, "SF Pro Text", sans-serif;
                font-size: 12px;
                font-weight: 400;
                padding: 0px;
                margin: 0px;
            }
        """)
        palette = self._search.palette()
        palette.setColor(QPalette.PlaceholderText, QColor("#48484A"))
        self._search.setPalette(palette)
        self._search.textChanged.connect(self._populate_leads)
        
        sc_layout.addWidget(s_icon, 0, Qt.AlignVCenter)
        sc_layout.addWidget(self._search, 1, Qt.AlignVCenter)
        search_lyt.addWidget(search_container)
        layout.addWidget(search_w)

        # C. Segmented Control
        seg_w = QWidget()
        seg_lyt = QVBoxLayout(seg_w)
        seg_lyt.setContentsMargins(10, 0, 10, 8)
        
        seg_track = QWidget()
        seg_track.setFixedHeight(30)
        seg_track.setStyleSheet("background-color: #2C2C2E; border-radius: 8px;")
        st_layout = QHBoxLayout(seg_track)
        st_layout.setContentsMargins(4, 4, 4, 4)
        st_layout.setSpacing(4)
        
        self._filter_btns = {}
        for label, key in [("All", "all"), ("Pending", "pending"), ("Sent", "sent")]:
            b = QPushButton(label.upper())
            b.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Preferred)
            b.setFixedHeight(22)
            b.setCursor(Qt.PointingHandCursor)
            st_layout.addWidget(b)
            self._filter_btns[key] = b
            b.clicked.connect(lambda _, k=key: self._on_filter(k))
            
        self._update_segment_styles("all")
        seg_lyt.addWidget(seg_track)
        layout.addWidget(seg_w)

        # D. Lead List
        self._lead_list = QListWidget()
        self._lead_list.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self._lead_list.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        self._lead_list.setFrameShape(QFrame.NoFrame)
        self._lead_list.setStyleSheet("""
            QListWidget { background: transparent; outline: none; border: none; }
            QListWidget::item { background: transparent; border: none; outline: none; }
            QListWidget::item:selected { background: transparent; border: none; outline: none; }
            QScrollBar:vertical { background: transparent; width: 3px; }
            QScrollBar::handle:vertical { background: #3A3A3C; border-radius: 2px; }
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical,
            QScrollBar::add-page:vertical, QScrollBar::sub-page:vertical { background: transparent; border: none; height: 0px; }
        """)
        self._lead_list.itemClicked.connect(self._on_lead_clicked)
        layout.addWidget(self._lead_list, 1)



        return left

    def _update_segment_styles(self, active_key: str):
        for key, btn in self._filter_btns.items():
            if key == active_key:
                btn.setStyleSheet("""
                    QPushButton {
                        background-color: #3A3A3C;
                        border-radius: 6px;
                        border: none;
                        font-family: 'PT Root UI', system-ui, sans-serif;
                        font-size: 10px;
                        color: #FFFFFF;
                        font-weight: 600;
                        padding: 2px 4px;
                    }
                """)
            else:
                btn.setStyleSheet("""
                    QPushButton {
                        background-color: transparent;
                        border: none;
                        font-family: 'PT Root UI', system-ui, sans-serif;
                        font-size: 10px;
                        color: #a0a0a5;
                        font-weight: 500;
                        padding: 2px 4px;
                    }
                    QPushButton:hover {
                        background: rgba(255,255,255,0.05);
                        border-radius: 6px;
                        color: #FFFFFF;
                    }
                """)

    def _build_center_panel(self) -> QWidget:
        center = QWidget()
        center.setStyleSheet("background: #141416;")
        layout = QVBoxLayout(center)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        # A. Letter Header
        header_w = QWidget()
        header_w.setFixedHeight(48)
        header_w.setStyleSheet("background-color: transparent; border: none;")
        h_layout = QHBoxLayout(header_w)
        h_layout.setContentsMargins(32, 0, 32, 10)
        h_layout.setSpacing(0)

        # Left
        left_h = QHBoxLayout()
        left_h.setSpacing(12)
        self._center_company = QLabel()
        self._center_company.setStyleSheet("color: #FFFFFF; font-family: 'PT Root UI'; font-size: 13px; font-weight: 600; border: none; background: transparent;")
        left_h.addWidget(self._center_company)

        def _top_info_label(icon_name):
            w = QWidget()
            w.setStyleSheet("background: transparent;")
            wl = QHBoxLayout(w)
            wl.setContentsMargins(0, 0, 0, 0)
            wl.setSpacing(4)
            ico = QLabel()
            ico.setFixedSize(13, 13)
            ico.setPixmap(_render_tinted_icon(icon_name, 13, "#8E8E93").pixmap(13, 13))
            ico.setStyleSheet("background: transparent;")
            lbl = QLabel()
            lbl.setStyleSheet("color: #8E8E93; font-family: 'PT Root UI'; font-size: 13px; font-weight: 400; border: none; background: transparent;")
            wl.addWidget(ico)
            wl.addWidget(lbl)
            return w, lbl

        hr_w, self._center_hr = _top_info_label("user.svg")
        job_w, self._center_job = _top_info_label("briefcase.svg")
        loc_w, self._center_city = _top_info_label("map-pin.svg")
        
        left_h.addWidget(hr_w)
        left_h.addWidget(job_w)
        left_h.addWidget(loc_w)
        h_layout.addLayout(left_h)

        h_layout.addStretch()

        # Right
        right_h = QHBoxLayout()
        right_h.setSpacing(8)
        
        def _header_icon_btn(fluent_icon, tooltip, is_danger=False):
            b = QPushButton()
            b.setFixedSize(32, 32)
            
            bg = "transparent"
            hover_bg = "rgba(255, 69, 58, 0.15)" if is_danger else "rgba(255, 255, 255, 0.08)"
            pressed_bg = "rgba(255, 69, 58, 0.25)" if is_danger else "rgba(255, 255, 255, 0.12)"
            color = "#FF453A" if is_danger else "#8E8E93"
            
            b.setIcon(fluent_icon.icon(color=color))
            b.setIconSize(QSize(16, 16))
            b.setCursor(Qt.PointingHandCursor)
            b.setToolTip(tooltip)
            b.setStyleSheet(f"""
                QPushButton {{
                    background: {bg};
                    border-radius: 6px;
                    border: none;
                }}
                QPushButton:hover {{ background: {hover_bg}; }}
                QPushButton:pressed {{ background: {pressed_bg}; }}
                QPushButton:disabled {{ background: transparent; }}
            """)
            return b

        self._btn_copy = _header_icon_btn(FluentIcon.COPY, "Copy letter")
        self._btn_polish = _header_icon_btn(FluentIcon.EDIT, "Polish letter")
        self._btn_regenerate = _header_icon_btn(FluentIcon.SYNC, "Regenerate letter")
        self._btn_undo_ver = _header_icon_btn(FluentIcon.HISTORY, "Revert to previous", is_danger=True)
        self._btn_undo_ver.setEnabled(False)

        self._btn_copy.clicked.connect(self._copy_letter)
        self._btn_polish.clicked.connect(self._polish_current_letter)
        self._btn_regenerate.clicked.connect(self._regenerate_current_letter)
        self._btn_undo_ver.clicked.connect(self._revert_to_previous)

        right_h.addWidget(self._btn_copy)
        right_h.addWidget(self._btn_polish)
        right_h.addWidget(self._btn_regenerate)
        right_h.addWidget(self._btn_undo_ver)
        h_layout.addLayout(right_h)
        
        layout.addWidget(header_w)

        # B. Letter Body
        self._editor = QTextEdit()
        self._editor.setFrameShape(QFrame.NoFrame)
        self._editor.setReadOnly(False)
        self._editor.setStyleSheet("""
            QTextEdit {
                background-color: transparent;
                border: none;
                padding: 20px 32px 32px 32px;
                color: #C7C7CC;
                font-family: 'PT Root UI', system-ui, -apple-system, sans-serif;
                font-size: 13px;
                font-weight: 400;
                line-height: 1.7;
            }
            QScrollBar:vertical { background: transparent; width: 3px; }
            QScrollBar::handle:vertical { background: #3A3A3C; border-radius: 2px; }
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical,
            QScrollBar::add-page:vertical, QScrollBar::sub-page:vertical { background: transparent; border: none; height: 0px; }
        """)
        self._editor.textChanged.connect(self._on_text_changed)
        
        # Find bar container
        self._find_bar = FindReplaceBar(self._editor, center)
        self._find_bar.hide()
        layout.addWidget(self._find_bar)

        layout.addWidget(self._editor, 1)

        return center

    def _create_sidebar_button(self, icon: FluentIcon, text: str, callback=None) -> QPushButton:
        # Add 3 spaces to pad the text away from the icon
        btn = QPushButton("   " + text.upper())
        btn.setIcon(icon.icon(color="#a0a0a5"))
        btn.setIconSize(QSize(14, 14))
        btn.setCursor(Qt.PointingHandCursor)
        btn.setStyleSheet("""
            QPushButton {
                background: transparent;
                border: none;
                border-radius: 6px;
                color: #a0a0a5;
                font-family: 'PT Root UI', sans-serif;
                font-size: 11px;
                font-weight: 600;
                letter-spacing: 1px;
                text-align: left;
                padding: 8px 12px;
            }
            QPushButton:hover {
                background: #2a2a2c;
                color: #FFFFFF;
            }
            QPushButton:disabled {
                color: rgba(160, 160, 165, 0.4);
            }
        """)
        if callback:
            btn.clicked.connect(callback)
        return btn


    def _build_right_panel(self) -> QScrollArea:
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        scroll.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        scroll.setFrameShape(QFrame.NoFrame)
        scroll.setStyleSheet("""
            QScrollArea { background: transparent; border: none; }
            QWidget#rightContainer { background: transparent; }
            QScrollBar:vertical { background: transparent; width: 3px; }
            QScrollBar::handle:vertical { background: #3A3A3C; border-radius: 2px; }
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical,
            QScrollBar::add-page:vertical, QScrollBar::sub-page:vertical { background: transparent; border: none; height: 0px; }
        """)

        container = QWidget()
        container.setObjectName("rightContainer")
        layout = QVBoxLayout(container)
        layout.setContentsMargins(8, 10, 8, 10)
        layout.setSpacing(16)

        def _create_card():
            card = QFrame()
            card.setObjectName("sidebarCard")
            card.setStyleSheet("QFrame#sidebarCard { background-color: #1a1a1c; border: 1px solid #2a2a2c; border-radius: 8px; }")
            cl = QVBoxLayout(card)
            cl.setContentsMargins(11, 10, 11, 10)
            cl.setSpacing(0)
            return card, cl

        def _create_divider():
            d = QFrame()
            d.setFixedHeight(1)
            d.setStyleSheet("background-color: #3A3A3C; max-height: 1px; border: none;")
            return d

        def _section_label(text):
            lbl = QLabel(text)
            lbl.setFixedHeight(16)
            lbl.setStyleSheet("color: #8E8E93; font-family: 'PT Root UI'; font-size: 11px; font-weight: 600; letter-spacing: 1.5px; background: transparent;")
            return lbl



        # CARD 4 - PROFILE & SIGNATURE
        card4, cl4 = _create_card()
        lbl4 = _section_label("PROFILE & SIGNATURE")
        cl4.addWidget(lbl4)
        cl4.addSpacing(7)
        self._btn_edit_profile = self._create_sidebar_button(FluentIcon.PEOPLE, "Set your information", self._edit_profile)
        self._btn_upload_sig = self._create_sidebar_button(FluentIcon.UPDATE, "Upload signature", self._upload_signature)
        self._btn_clear_sig = self._create_sidebar_button(FluentIcon.DELETE, "Clear signature", self._clear_signature)
        cl4.addWidget(self._btn_edit_profile)
        cl4.addSpacing(4)
        cl4.addWidget(self._btn_upload_sig)
        cl4.addSpacing(4)
        cl4.addWidget(self._btn_clear_sig)
        layout.addWidget(card4)

        # CARD 5 - TEMPLATE
        card5, cl5 = _create_card()
        lbl5 = _section_label("TEMPLATE")
        cl5.addWidget(lbl5)
        cl5.addSpacing(7)
        self._template_status = QLabel(self._template_status_text())
        self._template_status.setStyleSheet("color: #6E6E73; font-size: 11px; font-weight: 400; background: transparent;")
        cl5.addWidget(self._template_status)
        cl5.addSpacing(6)
        self._btn_edit_template = self._create_sidebar_button(FluentIcon.EDIT, "Edit template", self._edit_template)
        cl5.addWidget(self._btn_edit_template)
        layout.addWidget(card5)

        # CARD 6 - BEWERBUNG
        card6, cl6 = _create_card()
        lbl6 = _section_label("BEWERBUNG")
        cl6.addWidget(lbl6)
        cl6.addSpacing(7)

        self._bewerbung_status_label = QLabel("No PDF")
        self._bewerbung_status_label.setFixedHeight(22)
        self._bewerbung_status_label.setStyleSheet("""
            QLabel {
                background: #2C2C2E;
                border: 0.5px solid #3A3A3C;
                border-radius: 6px;
                color: #8E8E93;
                padding: 3px 10px;
                font-family: system-ui, -apple-system, "SF Pro Text", sans-serif;
                font-size: 11px;
                font-weight: 400;
            }
        """)
        
        status_row = QHBoxLayout()
        status_row.addWidget(self._bewerbung_status_label)
        status_row.addStretch()
        cl6.addLayout(status_row)
        cl6.addSpacing(6)

        # PAGE INDEX ROW (styled perfectly to match sidebar buttons)
        page_widget = QWidget()
        page_widget.setStyleSheet("QWidget { background: transparent; }")
        
        page_row = QHBoxLayout(page_widget)
        page_row.setContentsMargins(12, 4, 12, 4)
        page_row.setSpacing(0)
        
        from qfluentwidgets import IconWidget
        icon_w = IconWidget(FluentIcon.ALIGNMENT)
        icon_w.setFixedSize(14, 14)
        icon_w.setStyleSheet("background: transparent; color: #a0a0a5;")
        
        page_lbl = QLabel("   INSERT PAGE")
        page_lbl.setStyleSheet("""
            color: #a0a0a5;
            font-family: 'PT Root UI', sans-serif;
            font-size: 11px;
            font-weight: 600;
            letter-spacing: 1px;
            background: transparent;
            border: none;
        """)
        
        from PySide6.QtWidgets import QLineEdit
        self._page_input = QLineEdit()
        self._page_input.setFixedSize(28, 26)
        self._page_input.setAlignment(Qt.AlignCenter)
        self._page_input.setStyleSheet("""
            QLineEdit {
                background: rgba(0, 0, 0, 0.2);
                border: 1px solid #3A3A3C;
                border-radius: 4px;
                color: #FFFFFF;
                font-family: 'PT Root UI', sans-serif;
                font-size: 11px;
                font-weight: 600;
                padding: 0px;
                margin: 0px;
            }
            QLineEdit:focus {
                border: 1px solid #0A84FF;
                background: rgba(10, 132, 255, 0.1);
            }
        """)
        
        from ..core.config import config_manager
        saved_page = getattr(config_manager.settings, "bewerbung_anschreiben_page", 1)
        self._page_input.setText(str(saved_page))
        self._page_input.textChanged.connect(self._save_anschreiben_page)
        self._save_anschreiben_page()
        
        page_row.addWidget(icon_w)
        page_row.addWidget(page_lbl)
        page_row.addStretch()
        page_row.addWidget(self._page_input)
        
        cl6.addWidget(page_widget)
        cl6.addSpacing(4)

        def _show_export_menu():
            menu = RoundMenu(parent=self)
            menu.addAction(Action(FluentIcon.DOCUMENT, "Export .txt", triggered=self._export_letter))
            menu.addAction(Action(FluentIcon.DOCUMENT, "Export .docx", triggered=self._export_docx))
            menu.exec(self._btn_export.mapToGlobal(QPoint(0, self._btn_export.height())))

        self._btn_export = self._create_sidebar_button(FluentIcon.DOWNLOAD, "Export letter", _show_export_menu)
        self._btn_preview_pdf = self._create_sidebar_button(FluentIcon.VIEW, "Preview PDF", self._preview_merged_pdf)
        self._btn_add_page = self._create_sidebar_button(FluentIcon.ADD, "Load Bewerbung", self._on_pdf_browse)
        self._btn_close = self._create_sidebar_button(FluentIcon.CLOSE, "Clear PDF", self._clear_bewerbung_pdf)

        # Style Clear PDF button to be red
        self._btn_close.setIcon(FluentIcon.CLOSE.icon(color="#FF453A"))
        self._btn_close.setStyleSheet("""
            QPushButton {
                background: transparent;
                border: none;
                border-radius: 6px;
                color: #FF453A;
                font-family: 'PT Root UI', sans-serif;
                font-size: 11px;
                font-weight: 600;
                letter-spacing: 1px;
                text-align: left;
                padding: 8px 12px;
            }
            QPushButton:hover {
                background: rgba(255, 69, 58, 0.1);
            }
            QPushButton:disabled {
                color: rgba(255, 69, 58, 0.4);
            }
        """)

        cl6.addWidget(self._btn_add_page)
        cl6.addSpacing(4)
        cl6.addWidget(self._btn_preview_pdf)
        cl6.addSpacing(4)
        cl6.addWidget(self._btn_close)
        cl6.addSpacing(4)
        cl6.addWidget(self._btn_export)
        
        layout.addWidget(card6)

        # CARD 7 - BATCH ACTIONS
        card7, cl7 = _create_card()
        lbl7 = _section_label("BATCH ACTIONS")
        cl7.addWidget(lbl7)
        cl7.addSpacing(7)

        self._btn_sync_leads = self._create_sidebar_button(FluentIcon.DOWNLOAD, "Import leads", self._show_import_menu)
        self._btn_export_and_send_left = self._create_sidebar_button(FluentIcon.SEND, "Send all to queue", self._action_export_and_send_batch)
        self._btn_regen_all = self._create_sidebar_button(FluentIcon.SYNC, "Regenerate all", self._action_regenerate_all_letters)
        self._btn_delete_menu = self._create_sidebar_button(FluentIcon.DELETE, "Manage leads", self._show_delete_menu)
        
        def _style_btn(btn, icon, r, g, b):
            hex_color = f"#{r:02X}{g:02X}{b:02X}"
            btn.setIcon(icon.icon(color=hex_color))
            btn.setStyleSheet(f"""
                QPushButton {{
                    background: transparent;
                    border: none;
                    border-radius: 6px;
                    color: {hex_color};
                    font-family: 'PT Root UI', sans-serif;
                    font-size: 11px;
                    font-weight: 600;
                    letter-spacing: 1px;
                    text-align: left;
                    padding: 8px 12px;
                }}
                QPushButton:hover {{
                    background: rgba({r}, {g}, {b}, 0.1);
                }}
                QPushButton:disabled {{
                    color: rgba({r}, {g}, {b}, 0.4);
                }}
            """)

        _style_btn(self._btn_sync_leads, FluentIcon.DOWNLOAD, 10, 132, 255) # Blue
        _style_btn(self._btn_export_and_send_left, FluentIcon.SEND, 48, 209, 88) # Green
        _style_btn(self._btn_regen_all, FluentIcon.SYNC, 255, 159, 10) # Orange
        _style_btn(self._btn_delete_menu, FluentIcon.DELETE, 255, 69, 58) # Red

        cl7.addWidget(self._btn_sync_leads)
        cl7.addSpacing(4)
        cl7.addWidget(self._btn_export_and_send_left)
        cl7.addSpacing(4)
        cl7.addWidget(self._btn_regen_all)
        cl7.addSpacing(4)
        cl7.addWidget(self._btn_delete_menu)
        
        layout.addWidget(card7)

        layout.addStretch()
        scroll.setWidget(container)
        
        # Keep old references around so things don't crash when state is managed
        self._lead_summary = self._center_company
        self._status = self._center_company
        self._chip_sent = QLabel()
        self._chip_sent.hide()
        self._center_chip_sent = QLabel()
        self._center_chip_sent.hide()
        
        return scroll

    # ─────────────────────────────────────────────────────────────────────────
    # Keyboard shortcuts
    # ─────────────────────────────────────────────────────────────────────────

    def _setup_shortcuts(self):
        QShortcut(QKeySequence("Ctrl+H"), self).activated.connect(self._toggle_find_bar)
        QShortcut(QKeySequence("Ctrl+S"), self).activated.connect(self._autosave)
        QShortcut(QKeySequence("Escape"), self._find_bar).activated.connect(
            self._find_bar.hide
        )

    def _toggle_find_bar(self):
        if self._find_bar.isVisible():
            self._find_bar.hide()
            self._editor.setFocus()
        else:
            self._find_bar.show()
            self._find_bar.findChild(LineEdit).setFocus()
            self._find_bar.findChild(LineEdit).selectAll()

    def _toggle_find_bar(self):
        if self._find_bar.isVisible():
            self._find_bar.hide()
        else:
            self._find_bar.show_and_focus()

    # ─────────────────────────────────────────────────────────────────────────
    # Data loading
    # ─────────────────────────────────────────────────────────────────────────


    def _show_delete_menu(self):
        from PySide6.QtWidgets import QMenu
        from PySide6.QtGui import QAction
        
        menu = QMenu(self)
        menu.setWindowFlags(menu.windowFlags() | Qt.FramelessWindowHint | Qt.NoDropShadowWindowHint)
        menu.setAttribute(Qt.WA_TranslucentBackground)
        
        a1 = QAction("Delete selected", self)
        a1.triggered.connect(self._delete_selected)
        a2 = QAction("Delete sent", self)
        a2.triggered.connect(self._delete_sent)
        a3 = QAction("Delete all", self)
        a3.triggered.connect(self._delete_all)
        
        menu.addAction(a1)
        menu.addAction(a2)
        menu.addAction(a3)

        menu.setStyleSheet("""
            QMenu {
                background-color: #2C2C2E;
                border: 1px solid #3A3A3C;
                border-radius: 8px;
                padding: 6px;
            }
            QMenu::item {
                color: #FF453A;
                padding: 8px 32px 8px 16px;
                border-radius: 4px;
                margin: 2px 4px;
                font-family: 'PT Root UI', 'SF Pro Text', sans-serif;
                font-size: 13px;
                font-weight: 500;
            }
            QMenu::item:selected {
                background-color: rgba(255, 69, 58, 0.15);
            }
        """)
        menu.exec(self._btn_delete_menu.mapToGlobal(QPoint(0, self._btn_delete_menu.height() + 6)))

    def _delete_selected(self):
        if not self._selected_record:
            self._show_error("Delete Failed", "No lead selected.")
            return
            
        w = AppleConfirmDialog(
            "Delete Selected Lead",
            f"Are you sure you want to permanently delete {self._selected_record.company_name}?",
            parent=self.window()
        )
        if not w.exec():
            return
            
        import sqlite3
        try:
            db_path = get_memory_db_path()
            conn = sqlite3.connect(db_path, timeout=10.0)
            conn.execute("INSERT OR IGNORE INTO letter_state (lead_id) VALUES (?)", (self._selected_record.id,))
            conn.execute("UPDATE letter_state SET is_discarded = 1 WHERE lead_id = ?", (self._selected_record.id,))
            conn.commit()
            conn.close()
            
            self._lead_list.clearSelection()
            self._selected_record = None
            self._selected_state = None
            self.refresh()
            self._show_success("Discarded", "The lead has been removed from the Edit list.")
        except Exception as e:
            self._show_error("Delete Failed", str(e))

    def _delete_sent(self):
        w = AppleConfirmDialog(
            "Delete Sent Leads",
            "Are you sure you want to permanently delete all leads that have been marked as sent?",
            parent=self.window()
        )
        if not w.exec():
            return
            
        import sqlite3
        try:
            db_path = get_memory_db_path()
            conn = sqlite3.connect(db_path, timeout=10.0)
            # Find all lead IDs that are sent
            rows = conn.execute("SELECT lead_id FROM letter_state WHERE sent_at IS NOT NULL AND (is_discarded IS NULL OR is_discarded = 0)").fetchall()
            lead_ids = [r[0] for r in rows]
            
            if not lead_ids:
                self._show_success("Nothing to Remove", "There are no active sent leads.")
                conn.close()
                return
                
            placeholders = ",".join("?" * len(lead_ids))
            conn.execute(f"UPDATE letter_state SET is_discarded = 1 WHERE lead_id IN ({placeholders})", lead_ids)
            conn.commit()
            conn.close()
            
            self._lead_list.clearSelection()
            self._selected_record = None
            self._selected_state = None
            self.refresh()
            self._show_success("Discarded", f"Removed {len(lead_ids)} sent leads from the Edit list.")
        except Exception as e:
            self._show_error("Delete Failed", str(e))

    def _delete_all(self):
        w = AppleConfirmDialog(
            "Delete ALL Leads",
            "Are you absolutely sure you want to permanently delete ALL leads in your database? This action cannot be undone.",
            parent=self.window()
        )
        if not w.exec():
            return
            
        import sqlite3
        try:
            db_path = get_memory_db_path()
            conn = sqlite3.connect(db_path, timeout=10.0)
            
            # For "Delete All", we insert/update all leads that are currently in self._records to be discarded
            records_ids = [r.id for r in self._records]
            if records_ids:
                for chunk in [records_ids[i:i + 500] for i in range(0, len(records_ids), 500)]:
                    for rid in chunk:
                        conn.execute("INSERT OR IGNORE INTO letter_state (lead_id) VALUES (?)", (rid,))
                    placeholders = ",".join("?" * len(chunk))
                    conn.execute(f"UPDATE letter_state SET is_discarded = 1 WHERE lead_id IN ({placeholders})", chunk)
                    
            conn.commit()
            conn.close()
            
            self._lead_list.clearSelection()
            self._selected_record = None
            self._selected_state = None
            self.refresh()
            self._show_success("List Cleared", "All leads have been removed from the Edit list.")
        except Exception as e:
            self._show_error("Delete Failed", str(e))

    def refresh(self):
        self._status.setText("Loading leads…")
        run_in_thread(
            self._load_data,
            on_result=self._on_data_loaded,
            on_error=lambda tb: self._show_error(
                "Load failed", tb.splitlines()[-1] if tb else "Unknown error"
            ),
        )

    def _load_data(self) -> tuple[list[LeadRecord], dict[str, LetterState]]:
        db_path = get_memory_db_path()
        if not db_path.exists():
            return [], {}
        _, records = ExportService().load_project(str(db_path))
        states     = self._load_letter_states(db_path)
        return records, states

    def _load_letter_states(self, db_path: Path) -> dict[str, LetterState]:
        conn = sqlite3.connect(db_path, timeout=30.0)
        try:
            self._ensure_letter_state_table(conn)
            rows = conn.execute(
                "SELECT lead_id, generated_at, edited_at, sent_at, "
                "letter_text, previous_text, last_saved_ts, is_discarded FROM letter_state"
            ).fetchall()
            return {
                row[0]: LetterState(
                    row[1] or "", row[2] or "", row[3] or "",
                    row[4] or "", row[5] or "", row[6] or "", row[7] or 0
                )
                for row in rows
            }
        finally:
            conn.close()

    @staticmethod
    def _ensure_letter_state_table(conn: sqlite3.Connection):
        conn.execute("""
            CREATE TABLE IF NOT EXISTS letter_state (
                lead_id       TEXT PRIMARY KEY,
                generated_at  TEXT,
                edited_at     TEXT,
                sent_at       TEXT,
                letter_text   TEXT,
                previous_text TEXT,
                last_saved_ts TEXT,
                is_discarded  INTEGER DEFAULT 0
            )
        """)
        # migrate older schema that lacked new columns
        existing = {r[1] for r in conn.execute("PRAGMA table_info(letter_state)")}
        for col, ctype, default in [("previous_text", "TEXT", "''"), ("last_saved_ts", "TEXT", "''"), ("is_discarded", "INTEGER", "0")]:
            if col not in existing:
                conn.execute(f"ALTER TABLE letter_state ADD COLUMN {col} {ctype} DEFAULT {default}")
        conn.commit()

    @staticmethod
    def _ensure_settings_table(conn: sqlite3.Connection):
        conn.execute("""
            CREATE TABLE IF NOT EXISTS settings (
                key TEXT PRIMARY KEY,
                value TEXT
            )
        """)
        conn.commit()

    def _on_db_updated(self, records: list):
        run_in_thread(
            self._load_data,
            on_result=self._on_data_loaded
        )

    def _on_live_result_added(self, record):
        # ensure no duplicate
        for r in self._records:
            if r.id == record.id:
                return
        self._records.append(record)
        state = self._states.get(record.id)
        if state and state.is_discarded:
            return
            
        if self._active_filter == "sent" and not (state and state.sent_at):
            return
        if self._active_filter == "pending" and (state and state.sent_at):
            return
        
        needle = self._search.text().strip().lower() if hasattr(self, "_search") else ""
        if needle:
            haystack = " ".join(filter(None, [record.company_name, record.contact_person, record.city, record.job_title])).lower()
            if needle not in haystack:
                return

        item = QListWidgetItem()
        item.setData(Qt.UserRole, record.id)
        item.setSizeHint(QSize(0, 74))
        item.setToolTip(record.company_name or "Lead")
        self._lead_list.addItem(item)
        row_widget = LeadRowWidget(record, state)
        self._lead_list.setItemWidget(item, row_widget)
        self._update_filter_labels(needle)

    def _on_data_loaded(self, payload):
        self._records, self._states = payload
        self._populate_leads()
        if self._records:
            self._lead_list.setCurrentRow(0)
            self._render_record(self._records[0])
        else:
            self._editor.clear()
            self._status.setText("No leads found in SQLite.")
            self._set_actions_enabled(False)

    # ─────────────────────────────────────────────────────────────────────────
    # Lead list population
    # ─────────────────────────────────────────────────────────────────────────

    _active_filter: str = "all"

    def _on_filter(self, key: str):
        self._active_filter = key
        self._update_filter_labels()
        self._update_segment_styles(key)
        self._populate_leads()

    def _populate_leads(self):
        needle = self._search.text().strip().lower() if hasattr(self, "_search") else ""
        self._update_filter_labels(needle)
        self._lead_population_token += 1
        token = self._lead_population_token
        self._lead_list.blockSignals(True)
        self._lead_list.clear()
        self._lead_list.blockSignals(False)
        self._selected_row_widget = None

        self._pending_lead_records = []
        for record in self._records:
            state = self._states.get(record.id)

            if state and state.is_discarded:
                continue

            # apply filter
            if self._active_filter == "sent" and not (state and state.sent_at):
                continue
            if self._active_filter == "pending" and (state and state.sent_at):
                continue

            if needle:
                haystack = " ".join(
                    filter(None, [
                        record.company_name, record.contact_person,
                        record.city, record.job_title,
                    ])
                ).lower()
                if needle not in haystack:
                    continue

            self._pending_lead_records.append(record)

        if hasattr(self, "_sidebar_count"):
            pass

        self._pending_lead_index = 0
        QTimer.singleShot(0, lambda: self._append_lead_batch(token))

    def _update_filter_labels(self, needle: str | None = None) -> None:
        if not hasattr(self, "_filter_btns"):
            return
        if needle is None:
            needle = self._search.text().strip().lower() if hasattr(self, "_search") else ""

        counts = {"all": 0, "pending": 0, "sent": 0}
        for record in self._records:
            state = self._states.get(record.id)
            if state and state.is_discarded:
                continue
                
            if needle:
                haystack = " ".join(
                    filter(None, [
                        record.company_name, record.contact_person,
                        record.city, record.job_title,
                    ])
                ).lower()
                if needle not in haystack:
                    continue

            is_sent = bool(state and state.sent_at)
            counts["all"] += 1
            counts["sent" if is_sent else "pending"] += 1

        for key, btn in self._filter_btns.items():
            base_label = {"all": "ALL", "pending": "PENDING", "sent": "SENT"}[key]
            btn.setText(f"{base_label} {counts[key]}")

    def _append_lead_batch(self, token: int):
        if token != self._lead_population_token:
            return
        batch_size = 15
        end = min(self._pending_lead_index + batch_size, len(self._pending_lead_records))
        self._lead_list.setUpdatesEnabled(False)
        try:
            for record in self._pending_lead_records[self._pending_lead_index:end]:
                item = QListWidgetItem()
                item.setData(Qt.UserRole, record.id)
                item.setSizeHint(QSize(0, 74))
                item.setToolTip(record.company_name or "Lead")
                self._lead_list.addItem(item)
                row_widget = LeadRowWidget(record, self._states.get(record.id))
                row_widget.set_selected(record.id == self._selected_lead_id)
                if record.id == self._selected_lead_id:
                    self._selected_row_widget = row_widget
                    self._lead_list.setCurrentItem(item)
                self._lead_list.setItemWidget(item, row_widget)
        finally:
            self._lead_list.setUpdatesEnabled(True)

        self._pending_lead_index = end
        if self._pending_lead_index < len(self._pending_lead_records):
            QTimer.singleShot(15, lambda: self._append_lead_batch(token))

    # ─────────────────────────────────────────────────────────────────────────
    # Lead selection / rendering
    # ─────────────────────────────────────────────────────────────────────────

    def _on_lead_clicked(self, item: QListWidgetItem):
        lead_id = item.data(Qt.UserRole)
        record  = next((r for r in self._records if r.id == lead_id), None)
        if record:
            self._render_record(record)

    def _render_record(self, record: LeadRecord):
        # auto-save outgoing lead before switching
        if self._selected_record and not self._loading_text:
            self._autosave(silent=True)

        self._selected_record  = record
        previous_selected      = self._selected_row_widget
        self._selected_lead_id = record.id
        state  = self._states.get(record.id, LetterState())
        text   = state.letter_text or self._assemble_letter(record)

        self._rendered_text = text
        self._set_editor_text(text)

        self._set_actions_enabled(True)
        self._set_status(record, state)
        self._update_lead_summary(record)
        

        self._update_lead_selection_styles(previous_selected)
        self._update_word_count()
        self._check_placeholders()

        if not state.generated_at:
            self._mark_generated_async(record.id, text)

        if state.last_saved_ts:
            pass
        else:
            pass

        # enable/disable revert button
        self._set_undo_ver_enabled(bool(state.previous_text))

    def _update_lead_selection_styles(
        self, previous_selected: LeadRowWidget | None = None
    ):
        if previous_selected is not None:
            previous_selected.set_selected(False)
        for row in range(self._lead_list.count()):
            item   = self._lead_list.item(row)
            widget = self._lead_list.itemWidget(item)
            if isinstance(widget, LeadRowWidget) and item.data(Qt.UserRole) == self._selected_lead_id:
                widget.set_selected(True)
                self._selected_row_widget = widget
                break


    # ─────────────────────────────────────────────────────────────────────────
    # Auto-save
    # ─────────────────────────────────────────────────────────────────────────

    def _autosave(self, silent: bool = False):
        if not self._selected_record:
            return
        text  = self._editor.toPlainText()
        state = self._states.setdefault(self._selected_record.id, LetterState())
        if text == state.letter_text:
            return   # nothing changed
        now_ts = datetime.now().strftime("%H:%M")
        state.previous_text  = state.letter_text   # keep one version back
        state.letter_text    = text
        state.edited_at      = state.edited_at or datetime.utcnow().isoformat()
        state.last_saved_ts  = now_ts
        self._persist_state_async(
            self._selected_record.id,
            edited_at=state.edited_at,
            letter_text=text,
            previous_text=state.previous_text,
            last_saved_ts=now_ts,
        )
        self._set_undo_ver_enabled(bool(state.previous_text))

    # ─────────────────────────────────────────────────────────────────────────
    # Word count + placeholder audit
    # ─────────────────────────────────────────────────────────────────────────

    def _update_word_count(self):
        return

    def _check_placeholders(self):
        _highlight_placeholders(self._editor)

    # ─────────────────────────────────────────────────────────────────────────
    # Text-changed handler
    # ─────────────────────────────────────────────────────────────────────────

    def _set_editor_text(self, text: str):
        self._loading_text = True
        try:
            self._editor.setPlainText(text)
        finally:
            self._loading_text = False

    def _on_text_changed(self):
        if self._loading_text or not self._selected_record:
            return
        current = self._editor.toPlainText()
        state   = self._states.setdefault(self._selected_record.id, LetterState())
        if current != self._rendered_text:
            state.edited_at = state.edited_at or datetime.utcnow().isoformat()
        self._set_status(self._selected_record, state)
        self._autosave_timer.start()
        self._wc_timer.start()

    # ─────────────────────────────────────────────────────────────────────────
    # Letter assembly & polish
    # ─────────────────────────────────────────────────────────────────────────

    def _assemble_letter(self, record: LeadRecord) -> str:
        template = self._template_text or self._load_template()
        sender   = self._load_sender_settings()
        replacements = {
            "ANREDE":         self._salutation(record.contact_person),
            "FIRMA":          record.company_name        or "Unternehmen",
            "ORT":            record.city                or "",
            "PLZ":            record.postal_code         or "",
            "BERUF":          sender.get("beruf") or record.job_title or record.category or "Ausbildung",
            "DATUM":          self._german_date(),
        }
        for k, v in sender.items():
            replacements[f"SENDER_{k.upper()}"] = v
            
        text = template
        for key, value in replacements.items():
            text = text.replace(f"{{{{{key}}}}}", value)
        return self._polish_letter(text, replacements)

    def _fetch_sender_settings(self, **kwargs) -> None:
        run_in_thread(
            self._do_fetch_sender_settings,
            on_result=self._on_sender_settings_loaded
        )
        
    def _do_fetch_sender_settings(self) -> dict:
        import sqlite3
        try:
            db_path = get_memory_db_path()
            conn    = sqlite3.connect(db_path, timeout=10.0)
            self._ensure_settings_table(conn)
            rows    = conn.execute(
                "SELECT key, value FROM settings WHERE key LIKE 'sender_%'"
            ).fetchall()
            conn.close()
            return {row[0].replace("sender_", ""): row[1] for row in rows}
        except Exception:
            return {}

    def _on_sender_settings_loaded(self, settings: dict) -> None:
        self._cached_sender_settings = settings

    def _load_sender_settings(self) -> dict:
        """Return cached sender settings."""
        return self._cached_sender_settings.copy()

    def _do_fetch_persisted_pdf_path(self) -> str:
        import sqlite3
        try:
            db_path = get_memory_db_path()
            conn = sqlite3.connect(db_path, timeout=10.0)
            self._ensure_settings_table(conn)
            row = conn.execute("SELECT value FROM settings WHERE key = 'bewerbung_pdf_path'").fetchone()
            conn.close()
            return row[0] if row else ""
        except Exception:
            return ""

    def _on_persisted_pdf_path_loaded(self, path_str: str) -> None:
        if path_str:
            p = Path(path_str)
            if p.exists():
                self._load_bewerbung_pdf(p)

    def _do_fetch_signature_path(self) -> str:
        import sqlite3
        try:
            db_path = get_memory_db_path()
            conn = sqlite3.connect(db_path, timeout=10.0)
            self._ensure_settings_table(conn)
            row = conn.execute("SELECT value FROM settings WHERE key = 'signature_image_path'").fetchone()
            conn.close()
            return row[0] if row else ""
        except Exception:
            return ""

    def _on_signature_path_loaded(self, path_str: str) -> None:
        self._signature_image_path = path_str

    def _save_signature_path(self, path_str: str) -> None:
        self._signature_image_path = path_str
        def _do_save():
            import sqlite3
            try:
                db_path = get_memory_db_path()
                conn = sqlite3.connect(db_path, timeout=10.0)
                self._ensure_settings_table(conn)
                if path_str:
                    conn.execute("INSERT OR REPLACE INTO settings (key, value) VALUES ('signature_image_path', ?)", (path_str,))
                else:
                    conn.execute("DELETE FROM settings WHERE key = 'signature_image_path'")
                conn.commit()
                conn.close()
            except Exception:
                pass
        run_in_thread(_do_save)

    def _edit_profile(self) -> None:
        from PySide6.QtWidgets import QDialog, QVBoxLayout, QHBoxLayout, QLabel, QScrollArea, QWidget, QLineEdit, QInputDialog
        from qfluentwidgets import PushButton, PrimaryPushButton
        
        dialog = QDialog(self)
        dialog.setWindowTitle("Set Your Information")
        dialog.setMinimumSize(500, 600)
        dialog.setStyleSheet("QDialog { background: #1C1C1E; }")
        
        layout = QVBoxLayout(dialog)
        layout.setContentsMargins(24, 24, 24, 24)
        layout.setSpacing(16)
        
        # Header
        header = QLabel("Personal Information")
        header.setStyleSheet("color: #FFFFFF; font-family: system-ui, -apple-system, sans-serif; font-size: 18px; font-weight: 600; background: transparent;")
        layout.addWidget(header)
        
        hint = QLabel("Enter your details to automatically inject them into your templates.")
        hint.setWordWrap(True)
        hint.setStyleSheet("color: #8E8E93; font-family: system-ui, -apple-system, sans-serif; font-size: 12px; background: transparent;")
        layout.addWidget(hint)
        layout.addSpacing(10)
        
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setStyleSheet("QScrollArea { border: none; background: transparent; }")
        
        container = QWidget()
        container.setStyleSheet("background: transparent;")
        form_layout = QVBoxLayout(container)
        form_layout.setContentsMargins(0, 0, 16, 0)
        form_layout.setSpacing(16)
        scroll.setWidget(container)
        layout.addWidget(scroll, 1)
        
        current = self._cached_sender_settings.copy()
        fields = {}
        
        def _add_field(key, label_title, token_text):
            row = QVBoxLayout()
            row.setSpacing(6)
            
            lbl_row = QHBoxLayout()
            lbl_row.setContentsMargins(0, 0, 0, 0)
            
            lbl = QLabel(label_title.upper())
            lbl.setStyleSheet("color: #8E8E93; font-family: system-ui, -apple-system, sans-serif; font-weight: 600; font-size: 10px; letter-spacing: 1px; background: transparent;")
            
            token = QLabel(token_text)
            token.setStyleSheet("""
                QLabel {
                    background: rgba(10, 132, 255, 0.1);
                    color: #0A84FF;
                    border-radius: 4px;
                    padding: 2px 6px;
                    font-family: 'Menlo', monospace;
                    font-size: 10px;
                    font-weight: bold;
                }
            """)
            
            lbl_row.addWidget(lbl)
            lbl_row.addStretch(1)
            lbl_row.addWidget(token)
            row.addLayout(lbl_row)
            
            inp = QLineEdit()
            inp.setText(current.get(key, ""))
            inp.setStyleSheet("""
                QLineEdit {
                    background: rgba(255, 255, 255, 0.03); 
                    border: 1px solid rgba(255, 255, 255, 0.08);
                    border-radius: 6px; 
                    color: #E5E5EA; 
                    padding: 8px 12px;
                    font-size: 13px;
                }
                QLineEdit:focus { 
                    border: 1px solid #0A84FF; 
                    background: rgba(255, 255, 255, 0.05);
                }
            """)
            row.addWidget(inp)
            form_layout.addLayout(row)
            fields[key] = inp

        _add_field("name", "Full Name", "{{SENDER_NAME}}")
        _add_field("email", "Email Address", "{{SENDER_EMAIL}}")
        _add_field("phone", "Phone Number", "{{SENDER_PHONE}}")
        _add_field("address", "Street & House Number", "{{SENDER_ADDRESS}}")
        _add_field("city", "Postal Code & City", "{{SENDER_CITY}}")
        _add_field("beruf", "Target Job (Overrides Default)", "{{BERUF}}")
        
        default_keys = {"name", "email", "phone", "address", "city", "beruf"}
        for k in current.keys():
            if k not in default_keys:
                _add_field(k, f"Custom: {k}", f"{{{{SENDER_{k.upper()}}}}}")
                
        btn_add_custom = PushButton("+ ADD CUSTOM FIELD")
        btn_add_custom.setCursor(Qt.CursorShape.PointingHandCursor)
        btn_add_custom.setStyleSheet("""
            PushButton {
                background: transparent;
                border: 1px dashed rgba(255, 255, 255, 0.15);
                border-radius: 6px;
                color: #8E8E93;
                font-family: system-ui, -apple-system, sans-serif;
                font-size: 11px;
                font-weight: 600;
                letter-spacing: 1px;
                padding: 10px;
                min-height: 36px;
                max-height: 36px;
            }
            PushButton:hover {
                background: rgba(255, 255, 255, 0.03);
                border: 1px solid rgba(255, 255, 255, 0.25);
                color: #FFFFFF;
            }
        """)
        
        def _on_add_custom():
            key, ok = QInputDialog.getText(dialog, "New Custom Field", "Enter variable name:")
            if ok and key.strip():
                k = key.strip().lower()
                k = "".join(c for c in k if c.isalnum() or c == "_")
                if k and k not in fields:
                    _add_field(k, f"Custom: {k}", f"{{{{SENDER_{k.upper()}}}}}")
        btn_add_custom.clicked.connect(_on_add_custom)
        form_layout.addWidget(btn_add_custom)
        form_layout.addStretch()
        
        layout.addSpacing(8)
        btn_box = QHBoxLayout()
        btn_box.setSpacing(12)
        
        from PySide6.QtWidgets import QPushButton
        
        btn_cancel = QPushButton("Cancel")
        btn_cancel.setCursor(Qt.CursorShape.PointingHandCursor)
        btn_cancel.setFixedHeight(32)
        btn_cancel.setStyleSheet("""
            QPushButton {
                background: rgba(255, 255, 255, 0.1);
                border: none;
                border-radius: 6px;
                color: white;
                font-family: 'SF Pro Text', 'PT Root UI', sans-serif;
                font-size: 13px;
                font-weight: 500;
                padding: 0 16px;
            }
            QPushButton:hover { background: rgba(255, 255, 255, 0.15); }
        """)
        btn_cancel.clicked.connect(dialog.reject)
        
        btn_save = QPushButton("Save Information")
        btn_save.setCursor(Qt.CursorShape.PointingHandCursor)
        btn_save.setFixedHeight(32)
        btn_save.setStyleSheet("""
            QPushButton {
                background: #0A84FF;
                border: none;
                border-radius: 6px;
                color: #FFFFFF;
                font-family: 'SF Pro Text', 'PT Root UI', sans-serif;
                font-size: 13px;
                font-weight: 600;
                padding: 0 16px;
            }
            QPushButton:hover { background: #007AFF; }
        """)
        btn_save.clicked.connect(dialog.accept)
        
        btn_box.addStretch()
        btn_box.addWidget(btn_cancel)
        btn_box.addWidget(btn_save)
        layout.addLayout(btn_box)
        
        if dialog.exec():
            for k, inp in fields.items():
                self._cached_sender_settings[k] = inp.text().strip()
            
            def _do_save_all():
                import sqlite3
                try:
                    db_path = get_memory_db_path()
                    conn = sqlite3.connect(db_path, timeout=10.0)
                    self._ensure_settings_table(conn)
                    for k, v in self._cached_sender_settings.items():
                        key = f"sender_{k}"
                        if v:
                            conn.execute("INSERT OR REPLACE INTO settings (key, value) VALUES (?, ?)", (key, v))
                        else:
                            conn.execute("DELETE FROM settings WHERE key = ?", (key,))
                    conn.commit()
                    conn.close()
                except Exception:
                    pass
            run_in_thread(_do_save_all)
            self._show_success("Information Saved", "Your details have been updated for templates.")

    def _load_pdf_settings(self) -> dict:
        import sqlite3
        settings = {"font": "Helvetica", "size": "11", "leading": "14", "alignment": "Justified"}
        try:
            db_path = get_memory_db_path()
            conn = sqlite3.connect(db_path, timeout=10.0)
            self._ensure_settings_table(conn)
            rows = conn.execute("SELECT key, value FROM settings WHERE key LIKE 'pdf_%'").fetchall()
            for k, v in rows:
                settings[k.replace("pdf_", "")] = v
            conn.close()
        except Exception:
            pass
        return settings

    def _edit_pdf_settings(self) -> None:
        from PySide6.QtWidgets import QDialog, QVBoxLayout, QHBoxLayout, QLabel, QComboBox
        from qfluentwidgets import PushButton, PrimaryPushButton, ComboBox, SpinBox
        
        dialog = QDialog(self)
        dialog.setWindowTitle("PDF Styling Options")
        dialog.setMinimumWidth(350)
        dialog.setStyleSheet("QDialog { background: #1C1C1E; }")
        
        layout = QVBoxLayout(dialog)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(12)
        
        hint = QLabel("Customize the visual appearance of your generated application PDF.")
        hint.setWordWrap(True)
        hint.setStyleSheet("color: #AEAEB2; font-family: system-ui, -apple-system, sans-serif; font-size: 13px; background: transparent;")
        layout.addWidget(hint)
        layout.addSpacing(10)
        
        settings = self._load_pdf_settings()
        
        # Font Family
        lbl_font = QLabel("Font Family")
        lbl_font.setStyleSheet("color: #E5E5EA; font-weight: 500; font-size: 12px; background: transparent;")
        layout.addWidget(lbl_font)
        cb_font = ComboBox()
        cb_font.addItems(["Helvetica", "Times-Roman", "Courier"])
        cb_font.setCurrentText(settings.get("font", "Helvetica"))
        layout.addWidget(cb_font)
        
        # Font Size
        lbl_size = QLabel("Font Size (pt)")
        lbl_size.setStyleSheet("color: #E5E5EA; font-weight: 500; font-size: 12px; background: transparent;")
        layout.addWidget(lbl_size)
        sp_size = SpinBox()
        sp_size.setRange(8, 16)
        try:
            sp_size.setValue(int(settings.get("size", "11")))
        except ValueError:
            sp_size.setValue(11)
        layout.addWidget(sp_size)
        
        # Line Spacing
        lbl_lead = QLabel("Line Spacing (pt)")
        lbl_lead.setStyleSheet("color: #E5E5EA; font-weight: 500; font-size: 12px; background: transparent;")
        layout.addWidget(lbl_lead)
        sp_lead = SpinBox()
        sp_lead.setRange(10, 24)
        try:
            sp_lead.setValue(int(settings.get("leading", "14")))
        except ValueError:
            sp_lead.setValue(14)
        layout.addWidget(sp_lead)
        
        # Alignment
        lbl_align = QLabel("Text Alignment")
        lbl_align.setStyleSheet("color: #E5E5EA; font-weight: 500; font-size: 12px; background: transparent;")
        layout.addWidget(lbl_align)
        cb_align = ComboBox()
        cb_align.addItems(["Justified", "Left"])
        cb_align.setCurrentText(settings.get("alignment", "Justified"))
        layout.addWidget(cb_align)
        
        layout.addSpacing(16)
        btn_box = QHBoxLayout()
        btn_cancel = PushButton("Cancel")
        btn_cancel.clicked.connect(dialog.reject)
        btn_save = PrimaryPushButton("Save Styling")
        btn_save.clicked.connect(dialog.accept)
        btn_box.addStretch()
        btn_box.addWidget(btn_cancel)
        btn_box.addWidget(btn_save)
        layout.addLayout(btn_box)
        
        if dialog.exec():
            def _do_save_pdf_settings():
                import sqlite3
                try:
                    db_path = get_memory_db_path()
                    conn = sqlite3.connect(db_path, timeout=10.0)
                    self._ensure_settings_table(conn)
                    conn.execute("INSERT OR REPLACE INTO settings (key, value) VALUES ('pdf_font', ?)", (cb_font.currentText(),))
                    conn.execute("INSERT OR REPLACE INTO settings (key, value) VALUES ('pdf_size', ?)", (str(sp_size.value()),))
                    conn.execute("INSERT OR REPLACE INTO settings (key, value) VALUES ('pdf_leading', ?)", (str(sp_lead.value()),))
                    conn.execute("INSERT OR REPLACE INTO settings (key, value) VALUES ('pdf_alignment', ?)", (cb_align.currentText(),))
                    conn.commit()
                    conn.close()
                except Exception:
                    pass
            run_in_thread(_do_save_pdf_settings)
            self._show_success("Styling Saved", "Your PDF layout settings have been updated.")

    def _upload_signature(self) -> None:
        start_dir = QStandardPaths.writableLocation(QStandardPaths.PicturesLocation)
        if not start_dir:
            start_dir = str(Path.home())
        path, _ = QFileDialog.getOpenFileName(
            self,
            "Select Signature Image",
            start_dir,
            "Images (*.png *.jpg *.jpeg)"
        )
        if path:
            self._save_signature_path(path)
            self._show_success("Signature Saved", "Signature will be added to PDFs.")

    def _clear_signature(self) -> None:
        self._save_signature_path("")
        self._show_success("Signature Cleared", "Signature removed.")

    def _save_persisted_pdf_path(self, path_str: str) -> None:
        def _do_save():
            import sqlite3
            try:
                db_path = get_memory_db_path()
                conn = sqlite3.connect(db_path, timeout=10.0)
                self._ensure_settings_table(conn)
                if path_str:
                    conn.execute("INSERT OR REPLACE INTO settings (key, value) VALUES ('bewerbung_pdf_path', ?)", (path_str,))
                else:
                    conn.execute("DELETE FROM settings WHERE key = 'bewerbung_pdf_path'")
                conn.commit()
                conn.close()
            except Exception:
                pass
        run_in_thread(_do_save)

    def _polish_letter(self, text: str, replacements: dict[str, str]) -> str:
        cleaned = text.replace("\r\n", "\n").replace("\r", "\n").strip()
        cleaned = self._normalize_professional_german(cleaned)
        cleaned = re.sub(r"[ \t]+\n", "\n", cleaned)
        cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)

        company    = replacements.get("FIRMA")    or "Unternehmen"
        city       = replacements.get("ORT")      or ""
        job        = replacements.get("BERUF")    or "Ausbildung"
        salutation = replacements.get("ANREDE")   or "Sehr geehrte Damen und Herren,"
        date_line  = replacements.get("DATUM")    or self._german_date()

        if "{{" in cleaned or not self._looks_like_letter(cleaned):
            city_phrase = f" in {city}" if city else ""
            return "\n\n".join([
                f"{company}{city_phrase}",
                date_line,
                f"Bewerbung um einen Ausbildungsplatz als {job}",
                salutation,
                (
                    f"mit großem Interesse bewerbe ich mich bei {company} um einen "
                    f"Ausbildungsplatz als {job}. Die Möglichkeit, in einem professionellen "
                    "Umfeld praktische Erfahrung zu sammeln und mich fachlich "
                    "weiterzuentwickeln, spricht mich besonders an."
                ),
                (
                    "Ich arbeite zuverlässig, lerne schnell und gehe neue Aufgaben "
                    "sorgfältig und motiviert an. Teamfähigkeit, Verantwortungsbewusstsein "
                    "und eine klare Kommunikation sind für mich selbstverständlich."
                ),
                (
                    "Gerne überzeuge ich Sie in einem persönlichen Gespräch von meiner "
                    "Motivation. Über eine Einladung freue ich mich sehr."
                ),
                "Mit freundlichen Grüßen",
            ])

        # light lexical corrections
        corrections = {
            "mit grossem interesse":          "mit großem Interesse",
            "mit freundlichen gruessen":      "Mit freundlichen Grüßen",
            "ueber eine einladung":           "Über eine Einladung",
            "persoenlichen gespraech":        "persönlichen Gespräch",
            "zuverlaessig":                   "zuverlässig",
            "moechte":                        "möchte",
        }
        polished = cleaned
        for old, new in corrections.items():
            polished = re.sub(old, new, polished, flags=re.IGNORECASE)
        return polished

    def _normalize_professional_german(self, text: str) -> str:
        subs = {
            "grosses Interesse":              "großes Interesse",
            "grossem Interesse":              "großem Interesse",
            "grossen Interesse":              "großen Interesse",
            "Gruesse":                        "Grüße",
            "Mit freundlichen Gruessen":      "Mit freundlichen Grüßen",
            "persoenlichen Gespraech":        "persönlichen Gespräch",
            "persoenliches Gespraech":        "persönliches Gespräch",
            "zuverlaessig":                   "zuverlässig",
            "sorgfaeltig":                    "sorgfältig",
            "taetig":                         "tätig",
            "moechte":                        "möchte",
            "koennte":                        "könnte",
            "Ueber eine Einladung":           "Über eine Einladung",
        }
        normalized = text
        for old, new in subs.items():
            normalized = re.sub(re.escape(old), new, normalized, flags=re.IGNORECASE)
        return normalized

    def _looks_like_letter(self, text: str) -> bool:
        lower = text.lower()
        return "bewerbung" in lower and "mit freundlichen" in lower

    def _salutation(self, contact_person: str | None) -> str:
        name  = (contact_person or "").strip()
        if not name:
            return "Sehr geehrte Damen und Herren,"
        first = name.split()[0].strip(" ,.;:").lower()
        if first in {"frau", "ms", "mrs"}:
            clean = re.sub(r"^(frau|ms|mrs)\s+", "", name, flags=re.IGNORECASE).strip()
            return f"Sehr geehrte Frau {clean},"
        if first in {"herr", "mr"}:
            clean = re.sub(r"^(herr|mr)\s+", "", name, flags=re.IGNORECASE).strip()
            return f"Sehr geehrter Herr {clean},"
        return f"Guten Tag {name},"

    def _german_date(self) -> str:
        today = date.today()
        return f"{today.day:02d}. {GERMAN_MONTHS[today.month - 1]} {today.year}"

    # ─────────────────────────────────────────────────────────────────────────
    # Actions
    # ─────────────────────────────────────────────────────────────────────────

    def _copy_letter(self):
        if not self._selected_record:
            return
        QGuiApplication.clipboard().setText(self._editor.toPlainText())
        self._show_success("Copied", "Letter copied to clipboard.")

    def _export_letter(self):
        if not self._selected_record:
            return
        filename = self._letter_filename(self._selected_record, ".txt")
        path, _   = QFileDialog.getSaveFileName(
            self, "Export Anschreiben", filename, "Text Files (*.txt)"
        )
        if not path:
            return
        Path(path).write_text(self._editor.toPlainText(), encoding="utf-8")
        self._show_success("Exported", Path(path).name)

    def _export_docx(self):
        if not self._selected_record:
            return
        try:
            from docx import Document
            from docx.shared import Pt, Cm
        except ImportError:
            self._show_error(
                "python-docx missing",
                "Install it with:  pip install python-docx",
            )
            return

        filename = self._letter_filename(self._selected_record, ".docx")
        path, _   = QFileDialog.getSaveFileName(
            self, "Export as Word Document", filename, "Word Documents (*.docx)"
        )
        if not path:
            return

        doc = Document()
        for section in doc.sections:
            section.top_margin    = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin   = Cm(2.5)
            section.right_margin  = Cm(2.5)

        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        for line in self._editor.toPlainText().split("\n"):
            p   = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            p.paragraph_format.space_after  = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            # bold the Betreff line
            if ("Bewerbung" in line or "Betreff" in line) and 10 < len(line.strip()) < 120:
                run.bold = True

        doc.save(path)
        self._show_success("Exported .docx", Path(path).name)

    def _show_import_menu(self):
        from qfluentwidgets import RoundMenu, Action
        from PySide6.QtCore import QPoint
        menu = RoundMenu(parent=self)
        
        action_all = Action(FluentIcon.DOWNLOAD, "Import All Results", self)
        action_all.triggered.connect(lambda: self._do_import("all"))
        menu.addAction(action_all)
        
        menu.addSeparator()
        
        action_ausbildung = Action(FluentIcon.EDUCATION, "From Ausbildung", self)
        action_ausbildung.triggered.connect(lambda: self._do_import("ausbildung"))
        menu.addAction(action_ausbildung)
        
        action_aubi = Action(FluentIcon.EDUCATION, "From AubiPlus", self)
        action_aubi.triggered.connect(lambda: self._do_import("aubi"))
        menu.addAction(action_aubi)
        
        action_jobsuche = Action(FluentIcon.SEARCH, "From Jobsuche", self)
        action_jobsuche.triggered.connect(lambda: self._do_import("jobsuche"))
        menu.addAction(action_jobsuche)
        
        action_maps = Action(FluentIcon.PIN, "From Google Maps", self)
        action_maps.triggered.connect(lambda: self._do_import("maps"))
        menu.addAction(action_maps)
        
        menu.addSeparator()
        
        action_city = Action(FluentIcon.MARKET, "From Latest City", self)
        action_city.triggered.connect(lambda: self._do_import("city"))
        menu.addAction(action_city)
        
        menu.addSeparator()
        
        action_latest = Action(FluentIcon.SYNC, "From Latest Search", self)
        action_latest.triggered.connect(lambda: self._do_import("latest"))
        menu.addAction(action_latest)
        
        action_last_hour = Action(FluentIcon.HISTORY, "Last Hour", self)
        action_last_hour.triggered.connect(lambda: self._do_import("last_hour"))
        menu.addAction(action_last_hour)
        
        pos = self._btn_sync_leads.mapToGlobal(QPoint(self._btn_sync_leads.width() + 5, 0))
        menu.exec(pos)

    def _do_import(self, filter_type: str):
        from src.services.orchestrator import orchestrator
        all_records = orchestrator.get_app_memory_records()
        
        filtered = []
        if filter_type == "all":
            filtered = all_records
        elif filter_type == "maps":
            filtered = [r for r in all_records if r.source_type and "maps" in str(r.source_type).lower()]
        elif filter_type == "jobsuche":
            filtered = [r for r in all_records if r.source_type and "jobsuche" in str(r.source_type).lower()]
        elif filter_type == "ausbildung":
            filtered = [r for r in all_records if r.source_type and "ausbildung" in str(r.source_type).lower()]
        elif filter_type == "aubi":
            filtered = [r for r in all_records if r.source_type and "aubi" in str(r.source_type).lower()]
        elif filter_type == "city":
            if all_records:
                latest_city = all_records[-1].city
                if latest_city:
                    filtered = [r for r in all_records if r.city and r.city.lower() == latest_city.lower()]
        elif filter_type == "latest":
            if orchestrator.current_job and orchestrator.current_job.results:
                filtered = orchestrator.current_job.results
        elif filter_type == "last_hour":
            from datetime import datetime, timezone
            now = datetime.now(timezone.utc)
            for r in all_records:
                try:
                    dt = r.scraped_at
                    if isinstance(dt, str):
                        dt = datetime.fromisoformat(dt.replace("Z", "+00:00"))
                    if dt.tzinfo is None:
                        dt = dt.replace(tzinfo=timezone.utc)
                    if (now - dt).total_seconds() <= 3600:
                        filtered.append(r)
                except: pass
                
        if not filtered:
            self._show_error("Import Failed", "No leads match the selected criteria.")
            return
            
        def _persist():
            from src.core.config import get_memory_db_path
            from src.core.models import ScrapingJob, SearchConfig, ScrapingStatus
            from src.services.export_service import ExportService
            import sqlite3
            
            db_path = str(get_memory_db_path())
            dummy_job = ScrapingJob(config=SearchConfig(), results=filtered, status=ScrapingStatus.COMPLETED)
            ExportService().save_project(dummy_job, db_path)
            
            conn = sqlite3.connect(db_path, timeout=10.0)
            try:
                records_ids = [r.stable_id() for r in filtered]
                for chunk in [records_ids[i:i + 500] for i in range(0, len(records_ids), 500)]:
                    placeholders = ",".join("?" * len(chunk))
                    conn.execute(f"UPDATE letter_state SET is_discarded = 0 WHERE lead_id IN ({placeholders})", chunk)
                conn.commit()
            finally:
                conn.close()

        def _on_persist_done(*args):
            self.refresh()
            self._show_success("Imported", f"{len(filtered)} search results have been imported.")

        from src.utils.db_worker import run_in_thread
        run_in_thread(
            _persist,
            on_finished=_on_persist_done
        )

    def _batch_export(self):
        """Export letters for all currently visible (filtered) leads."""
        if not self._pending_lead_records:
            self._show_error("Nothing to export", "No leads match current filter.")
            return

        folder = QFileDialog.getExistingDirectory(
            self, "Choose export folder", ""
        )
        if not folder:
            return

        folder_path = Path(folder)
        count       = 0
        for record in self._pending_lead_records:
            state = self._states.get(record.id, LetterState())
            text  = state.letter_text or self._assemble_letter(record)
            fname = self._letter_filename(record, ".txt")
            (folder_path / fname).write_text(text, encoding="utf-8")
            count += 1

        self._show_success("Batch Export", f"{count} letters saved to {folder_path.name}/")



    def _toggle_preview(self):
        pass

    def _revert_to_previous(self):
        if not self._selected_record:
            return
        state = self._states.get(self._selected_record.id)
        if not state or not state.previous_text:
            self._show_error("No previous version", "Nothing to revert to.")
            return
        dialog = AppleConfirmDialog(
            "Revert",
            "Replace current letter with the previous saved version?",
            "REVERT",
            "#FF453A",
            self.window()
        )
        if not dialog.exec():
            return
        self._set_editor_text(state.previous_text)
        state.letter_text  = state.previous_text
        state.previous_text = ""
        self._set_undo_ver_enabled(False)
        self._show_success("Reverted", "Previous version restored.")

    def _polish_current_letter(self):
        if not self._selected_record:
            return
        sender = self._load_sender_settings()
        replacements = {
            "ANREDE": self._salutation(self._selected_record.contact_person),
            "FIRMA":  self._selected_record.company_name or "Unternehmen",
            "ORT":    self._selected_record.city or "",
            "PLZ":    self._selected_record.postal_code or "",
            "BERUF":  sender.get("beruf") or self._selected_record.job_title or self._selected_record.category or "Ausbildung",
            "DATUM":  self._german_date(),
        }
        polished = self._polish_letter(self._editor.toPlainText(), replacements)
        self._set_editor_text(polished)
        state = self._states.setdefault(self._selected_record.id, LetterState())
        state.edited_at   = datetime.utcnow().isoformat()
        state.letter_text = polished
        self._set_status(self._selected_record, state)
        self._persist_state_async(
            self._selected_record.id,
            edited_at=state.edited_at,
            letter_text=polished,
        )
        self._check_placeholders()
        self._show_success("Polished", "Letter wording has been refined.")

    def _action_regenerate_all_letters(self):
        w = AppleConfirmDialog(
            "Regenerate All",
            "Are you sure you want to regenerate all visible letters? Any manual edits will be overwritten.",
            confirm_text="Regenerate",
            confirm_color="#0A84FF",
            parent=self
        )
        if w.exec() != QDialog.Accepted:
            return
            
        total = len(self._pending_lead_records)
        import sqlite3
        from ..core.config import get_memory_db_path
        try:
            conn = sqlite3.connect(str(get_memory_db_path()), timeout=10.0)
            for record in self._pending_lead_records:
                text = self._assemble_letter(record)
                state = self._states.setdefault(record.id, LetterState())
                state.letter_text = text
                state.edited_at   = ""
                conn.execute(
                    "INSERT INTO letter_state (lead_id, letter_text, edited_at) VALUES (?, ?, '') "
                    "ON CONFLICT(lead_id) DO UPDATE SET letter_text=excluded.letter_text, edited_at=''",
                    (record.id, text)
                )
            conn.commit()
            conn.close()
            
            if self._selected_record:
                self._regenerate_current_letter()
            
            self._show_success("Success", f"Regenerated {total} letters.")
            
            from ..core.events import event_bus, EventBus
            event_bus.emit(EventBus.DB_UPDATED, records=[])
        except Exception as e:
            self._show_error("Error", f"Failed to regenerate all: {e}")

    def _regenerate_current_letter(self):
        if not self._selected_record:
            return
        text = self._assemble_letter(self._selected_record)
        self._rendered_text = text
        self._set_editor_text(text)
        state = self._states.setdefault(self._selected_record.id, LetterState())
        state.letter_text = text
        state.edited_at   = ""
        self._set_status(self._selected_record, state)
        self._persist_state_async(self._selected_record.id, letter_text=text)
        self._check_placeholders()

    # ─────────────────────────────────────────────────────────────────────────
    # Template management
    # ─────────────────────────────────────────────────────────────────────────



    def _edit_template(self):
        dialog = TemplateEditorDialog(self, self._template_text or self._load_template())
        if dialog.exec():
            self._save_template(dialog.get_template_text())
            self._show_success("Template Saved", "Regenerating selected lead.")
            self._regenerate_current_letter()

    def _load_template(self) -> str:
        if self._template_path.exists():
            return self._template_path.read_text(encoding="utf-8")
        return (
            "{{SENDER_NAME}}\n{{SENDER_ADDRESS}}\n{{SENDER_CITY}}\n"
            "{{SENDER_PHONE}} · {{SENDER_EMAIL}}\n\n"
            "{{FIRMA}}\n{{PLZ}} {{ORT}}\n\n"
            "{{DATUM}}\n\n"
            "Bewerbung um einen Ausbildungsplatz als {{BERUF}}\n\n"
            "{{ANREDE}}\n\n"
            "mit großem Interesse bewerbe ich mich bei {{FIRMA}} um einen "
            "Ausbildungsplatz als {{BERUF}}.\n\n"
            "Mit freundlichen Grüßen"
        )

    def _save_template(self, text: str):
        self._template_path.parent.mkdir(parents=True, exist_ok=True)
        self._template_path.write_text(text.strip() + "\n", encoding="utf-8")
        self._template_text = text.strip() + "\n"
        self._template_status.setText(self._template_status_text())

    def _template_status_text(self) -> str:
        if self._template_path.exists():
            return f"Using: {self._template_path.name}"
        return "Using: built-in default"

    # ─────────────────────────────────────────────────────────────────────────
    # Persistence helpers
    # ─────────────────────────────────────────────────────────────────────────

    def _mark_generated_async(self, lead_id: str, text: str):
        now   = datetime.utcnow().isoformat()
        state = self._states.setdefault(lead_id, LetterState())
        state.generated_at = now
        state.letter_text  = text
        self._persist_state_async(lead_id, generated_at=now, letter_text=text)
        self._refresh_selected_lead_row()

    def _refresh_selected_lead_row(self):
        if not self._selected_record:
            return
        for row in range(self._lead_list.count()):
            item = self._lead_list.item(row)
            if item.data(Qt.UserRole) != self._selected_record.id:
                continue
            replacement = LeadRowWidget(
                self._selected_record,
                self._states.get(self._selected_record.id),
            )
            replacement.set_selected(True)
            self._lead_list.setItemWidget(item, replacement)
            self._selected_row_widget = replacement
            break

    def _persist_state_async(
        self,
        lead_id: str,
        *,
        generated_at:  str | None = None,
        edited_at:     str | None = None,
        sent_at:       str | None = None,
        letter_text:   str | None = None,
        previous_text: str | None = None,
        last_saved_ts: str | None = None,
    ):
        run_in_thread(
            self._persist_state,
            lead_id,
            generated_at,
            edited_at,
            sent_at,
            letter_text,
            previous_text,
            last_saved_ts,
            on_error=lambda tb: self._show_error(
                "Save failed", tb.splitlines()[-1] if tb else "Unknown error"
            ),
        )

    def _persist_state(
        self,
        lead_id:       str,
        generated_at:  str | None,
        edited_at:     str | None,
        sent_at:       str | None,
        letter_text:   str | None,
        previous_text: str | None,
        last_saved_ts: str | None,
    ):
        db_path = get_memory_db_path()
        conn    = sqlite3.connect(db_path, timeout=30.0)
        try:
            self._ensure_letter_state_table(conn)
            existing = conn.execute(
                "SELECT generated_at, edited_at, sent_at, letter_text, "
                "previous_text, last_saved_ts, is_discarded FROM letter_state WHERE lead_id=?",
                (lead_id,),
            ).fetchone()
            cur = LetterState(*(existing or ("", "", "", "", "", "", 0)))
            final = LetterState(
                generated_at  or cur.generated_at,
                edited_at     or cur.edited_at,
                sent_at       or cur.sent_at,
                letter_text   if letter_text   is not None else cur.letter_text,
                previous_text if previous_text is not None else cur.previous_text,
                last_saved_ts if last_saved_ts is not None else cur.last_saved_ts,
                cur.is_discarded,
            )
            conn.execute(
                """
                INSERT OR REPLACE INTO letter_state
                  (lead_id, generated_at, edited_at, sent_at,
                   letter_text, previous_text, last_saved_ts, is_discarded)
                VALUES (?,?,?,?,?,?,?,?)
                """,
                (
                    lead_id,
                    final.generated_at, final.edited_at, final.sent_at,
                    final.letter_text,  final.previous_text, final.last_saved_ts,
                    final.is_discarded,
                ),
            )
            conn.commit()
        finally:
            conn.close()

    # ─────────────────────────────────────────────────────────────────────────
    # UI helpers
    # ─────────────────────────────────────────────────────────────────────────

    def _letter_filename(self, record: LeadRecord, ext: str = ".txt") -> str:
        company = re.sub(
            r"[^A-Za-z0-9ÄÖÜäöüß_-]+", "_", record.company_name or "Firma"
        ).strip("_")
        stamp = date.today().strftime("%Y-%m-%d")
        return f"{company}_{stamp}{ext}"

    def _set_actions_enabled(self, enabled: bool):
        for btn in (
            self._btn_copy, self._btn_polish,
            self._btn_export,
            self._btn_undo_ver,
        ):
            btn.setEnabled(enabled)

    def _set_status(self, record: LeadRecord, state: LetterState):
        has_sent = bool(state.sent_at)

    def _update_lead_summary(self, record: LeadRecord):
        self._center_company.setText(record.company_name or "Unknown company")
        self._center_hr.setText(record.contact_person or "No HR contact")
        self._center_job.setText(record.job_title or record.category or "—")
        loc = " ".join(filter(None, [record.postal_code, record.city]))
        self._center_city.setText(loc or "No location")
        # keep legacy reference consistent
        self._lead_summary = self._center_company
        self._status = self._center_company

    # ─────────────────────────────────────────────────────────────────────────
    # Style factories
    # ─────────────────────────────────────────────────────────────────────────

    def _info_row_label(self, text: str) -> QLabel:
        lbl = QLabel(text)
        lbl.setStyleSheet(
            "color: #8E8E93; font-size: 10px; background: transparent; border: none;"
        )
        return lbl

    def _panel(self) -> QFrame:
        f = QFrame()
        f.setStyleSheet(
            "QFrame { background: transparent; border: none; }"
        )
        return f

    def _section_label(self, text: str) -> QLabel:
        lbl = QLabel(text)
        lbl.setStyleSheet(
            "color: #6E6E73; font-family: system-ui, -apple-system, sans-serif; "
            "font-size: 10px; font-weight: 400; letter-spacing: 0.06em; "
            "background: transparent; border: none;"
        )
        return lbl

    def _status_chip(self, text: str, color: str) -> QLabel:
        chip = QLabel(text)
        chip.setAlignment(Qt.AlignCenter)
        chip.setFixedHeight(20)
        bg = "rgba(48,209,88,0.18)" if color.lower() == "#30d158" else "rgba(10,132,255,0.18)"
        chip.setStyleSheet(
            f"color: {color}; background: {bg};"
            " border: none; border-radius: 5px;"
            " font-family: system-ui, -apple-system, sans-serif; font-size: 10px;"
            " font-weight: 500; padding: 0 7px;"
        )
        return chip

    def _pill_style(self, color: str) -> str:
        return (
            f"color: {color}; background: rgba(10,132,255,0.10); "
            f"border: 1px solid {color}; border-radius: 7px; "
            "font-family: system-ui, -apple-system, sans-serif; font-size: 10px; "
            "font-weight: 500; padding: 3px 8px;"
        )

    def _button(self, text: str, primary: bool = False) -> PushButton:
        btn = PushButton(text)
        btn.setFixedHeight(36)
        btn.setCursor(Qt.PointingHandCursor)
        btn.setStyleSheet(
            Theme.zugzwang_primary_button() if primary else Theme.zugzwang_button()
        )
        return btn

    def _small_button(self, text: str, primary: bool = False) -> PushButton:
        btn = PushButton(text)
        btn.setFixedHeight(30)
        btn.setCursor(Qt.PointingHandCursor)
        btn.setStyleSheet(
            Theme.zugzwang_primary_button() if primary else Theme.zugzwang_button()
        )
        return btn

    def _toolbar_button(self, text: str, color: str, tooltip: str, compact: bool = False) -> PushButton:
        btn = PushButton(text)
        size = 34 if compact else 42
        btn.setFixedSize(size, 34)
        btn.setCursor(Qt.PointingHandCursor)
        btn.setToolTip(tooltip)
        btn.setStyleSheet(f"""
            QPushButton {{
                background: {color};
                border: 1px solid rgba(255,255,255,0.14);
                border-radius: 9px;
                color: #FFFFFF;
                font-family: system-ui, -apple-system, sans-serif;
                font-size: 12px;
                font-weight: 500;
                padding: 0;
            }}
            QPushButton:hover {{
                background: rgba(255,255,255,0.12);
                border-color: {color};
                color: {color};
            }}
            QPushButton:pressed {{
                background: rgba(255,255,255,0.06);
            }}
            QPushButton:disabled {{
                background: rgba(255,255,255,0.06);
                border-color: rgba(255,255,255,0.06);
                color: rgba(255,255,255,0.3);
            }}
        """)
        return btn

    def _separator(self) -> QFrame:
        line = QFrame()
        line.setFrameShape(QFrame.HLine)
        line.setFixedHeight(1)
        line.setStyleSheet("background: rgba(255,255,255,0.06); border: none;")
        return line

    def _action_group(self, label: str) -> tuple:
        """Return (card QFrame, inner QVBoxLayout). Add _action_row widgets to the layout."""
        card = QFrame()
        card.setStyleSheet(
            "QFrame { background: transparent; border: none; }"
        )
        vbox = QVBoxLayout(card)
        vbox.setContentsMargins(0, 0, 0, 0)
        vbox.setSpacing(6)
        
        # Add spacing before section label
        vbox.addSpacing(20)
        
        lbl = QLabel(label)
        lbl.setFixedHeight(24)
        lbl.setStyleSheet("""
            color: #444444;
            font-family: system-ui, -apple-system, sans-serif;
            font-size: 9px;
            font-weight: 500;
            letter-spacing: 0.12em;
            background: transparent;
            border: none;
            border-bottom: 1px solid #1c1c1c;
            margin-bottom: 8px;
        """)
        vbox.addWidget(lbl)
        return card, vbox

    def _add_hairline(self, layout: QVBoxLayout):
        layout.addSpacing(4)

    def _set_undo_ver_enabled(self, enabled: bool):
        self._btn_undo_ver.setEnabled(enabled)
        if hasattr(self, "_btn_undo_ver_opacity"):
            self._btn_undo_ver_opacity.setOpacity(1.0 if enabled else 0.4)

    def _action_row(
        self,
        icon_name: str,
        text: str,
        style: str = "normal",   # normal | primary | danger
        icon_color: str | None = None,
        text_color: str | None = None,
        shortcut: str = "",
    ) -> QPushButton:
        """Full-width flat action row styled like the HTML mockup."""
        btn = QPushButton()
        btn.setFixedHeight(27)
        btn.setCursor(Qt.PointingHandCursor)
        btn.setFlat(True)

        if style == "danger" or "CLEAR" in text.upper():
            display_color = "#e05555"
            border_color = "#2a1414"
            hover_bg = "#3a1a1a"
            hover_border = "#4a1c1c"
            hover_fg = "#ff6b6b"
            bg = "#1a0c0c"
        else:
            display_color = text_color if text_color else "#AFAFAF"
            border_color = "#3A3A3C"
            hover_bg = "#48484A"
            hover_border = "#5C5C5E"
            hover_fg = "#FFFFFF"
            bg = "#2C2C2E"

        ico_col = icon_color if icon_color else display_color
        btn.setIcon(_render_tinted_icon(icon_name, 12, ico_col))
        btn.setIconSize(QSize(12, 12))
        btn.setText(text)

        btn.setStyleSheet(f"""
            QPushButton {{
                background: {bg};
                border: 1px solid {border_color};
                border-radius: 6px;
                color: {display_color};
                font-family: system-ui, -apple-system, sans-serif;
                font-size: 11px;
                font-weight: 500;
                padding-left: 8px;
                padding-right: 8px;
                text-align: left;
            }}
            QPushButton:hover {{
                background: {hover_bg};
                border-color: {hover_border};
                color: {hover_fg};
            }}
            QPushButton:pressed {{
                background: #1C1C1E;
            }}
            QPushButton:disabled {{
                background: transparent;
                color: #444444;
                border-color: #222222;
            }}
        """)
        return btn

    def _ghost_action_style(self, active: bool = False) -> str:
        if active:
            return """
                QPushButton {
                    background: #0A84FF;
                    border: none;
                    border-radius: 7px;
                    color: #FFFFFF;
                    font-family: system-ui, -apple-system, sans-serif;
                    font-size: 12px;
                    font-weight: 400;
                }
                QPushButton:hover { background: #1A8FFF; }
            """
        else:
            return """
                QPushButton {
                    background: #3A3A3C;
                    border: none;
                    border-radius: 7px;
                    color: #E5E5EA;
                    font-family: system-ui, -apple-system, sans-serif;
                    font-size: 12px;
                    font-weight: 400;
                }
                QPushButton:hover { background: #48484A; }
            """

    def _mappe_btn(
        self,
        icon_name: str,
        bg_color: str,
        tooltip: str,
        border: bool = False,
        icon_color: str = "#FFFFFF",
    ) -> QPushButton:
        return MappeButton(icon_name, tooltip, self)

    def _square_icon_button(self, text: str, color: str, tooltip: str) -> QPushButton:
        btn = QPushButton(text)
        btn.setFixedSize(44, 44)
        btn.setCursor(Qt.PointingHandCursor)
        btn.setToolTip(tooltip)
        border = color
        btn.setStyleSheet(f"""
            QPushButton {{
                background: {color};
                border: 1px solid {border};
                border-radius: 9px;
                color: #FFFFFF;
                font-family: system-ui, -apple-system, sans-serif;
                font-size: 25px;
                font-weight: 400;
                padding: 0;
            }}
            QPushButton:hover {{
                background: rgba(255,255,255,0.12);
                border-color: rgba(255,255,255,0.18);
            }}
            QPushButton:pressed {{
                background: rgba(255,255,255,0.08);
            }}
            QPushButton:disabled {{
                background: #2C2C2E;
                border-color: #3A3A3C;
                color: #636366;
            }}
        """)
        return btn

    def _input_style(self) -> str:
        return """
            LineEdit, QLineEdit {
                background: #1C1C1E;
                border: 1px solid #3A3A3C;
                border-radius: 8px;
                color: #F2F2F7;
                padding: 0 10px;
            }
            LineEdit:focus, QLineEdit:focus { border-color: #0A84FF; }
        """

    def _list_style(self) -> str:
        return """
            QListWidget {
                background: transparent;
                border: none;
                color: #F2F2F7;
                outline: none;
                padding: 0;
            }
            QListWidget::item:selected { background: transparent; }
        """

    def _editor_style(self, preview: bool = False) -> str:
        return """
            QTextEdit {
                background: #0d0d0d;
                border: none;
                color: #c8c8c8;
                font-family: system-ui, -apple-system, sans-serif;
                font-size: 13px;
            }
        """

    def _filter_btn_style(self, active: bool) -> str:
        if active:
            return (
                "QPushButton { "
                "background: #FFFFFF; border: none; border-radius: 7px; "
                "color: #1C1C1E; font-family: system-ui, -apple-system, sans-serif; font-size: 11px; font-weight: 500; "
                "}"
            )
        return (
            "QPushButton { background: transparent; border: none; border-radius: 6px; "
            "color: #8E8E93; font-family: system-ui, -apple-system, sans-serif; font-size: 11px; font-weight: 500; }"
            "QPushButton:hover { color: #F2F2F7; background: rgba(255,255,255,0.05); }"
        )

    def _toggle_btn_style(self, active: bool) -> str:
        if active:
            return (
                "QPushButton { background: #0A84FF; "
                "border: 1px solid #0A84FF; border-radius: 8px; "
                "color: #FFFFFF; font-family: system-ui, -apple-system, sans-serif; "
                "font-size: 11px; font-weight: 500; }"
            )
        return (
            "QPushButton { background: rgba(255,255,255,0.04); "
            "border: 1px solid rgba(255,255,255,0.12); border-radius: 8px; "
            "color: #AEAEB2; font-family: system-ui, -apple-system, sans-serif; "
            "font-size: 11px; font-weight: 500; }"
            "QPushButton:hover { background: rgba(255,255,255,0.08); color: #F2F2F7; }"
        )

    # ─────────────────────────────────────────────────────────────────────────
    # Notifications
    # ─────────────────────────────────────────────────────────────────────────

    def _show_success(self, title: str, message: str):
        InfoBar.success(title, message, duration=2500, parent=self)

    def _show_error(self, title: str, message: str):
        InfoBar.error(title, message, duration=5000, parent=self)

    # ─────────────────────────────────────────────────────────────────────────
    # Bewerbungsmappe PDF Export Feature
    # ─────────────────────────────────────────────────────────────────────────

    def _build_bewerbungsmappe_section(self, layout: QVBoxLayout) -> None:
        # Outer card identical to other action-group cards
        card = QFrame()
        card.setStyleSheet(
            "QFrame { background: #1C1C1E; border: 1px solid rgba(255,255,255,0.07);"
            " border-radius: 9px; }"
        )
        vbox = QVBoxLayout(card)
        vbox.setContentsMargins(12, 10, 12, 10)
        vbox.setSpacing(6)

        lbl = QLabel("BEWERBUNGSMAPPE")
        lbl.setStyleSheet(
            "color: #3A3A3C; font-family: system-ui, -apple-system, sans-serif; font-size: 9px; font-weight: 500; letter-spacing: 0.1em;"
            " background: transparent; border: none;"
        )
        vbox.addWidget(lbl)

        # ── Drop zone (shown when no PDF loaded) ─────────────────────────────
        self._drop_zone = PDFDropZone(self)
        self._drop_zone.setFixedHeight(64)
        self._drop_zone.file_dropped.connect(self._on_pdf_dropped)
        self._drop_zone.clicked.connect(self._on_pdf_browse)
        vbox.addWidget(self._drop_zone)

        # ── Loaded state (shown after PDF upload) ─────────────────────────────
        self._file_info_widget = QFrame()
        self._file_info_widget.setStyleSheet(
            "QFrame { background: rgba(48,209,88,0.06); border: 1px solid rgba(48,209,88,0.2);"
            " border-radius: 8px; }"
        )
        self._file_info_widget.setVisible(False)
        fi_row = QHBoxLayout(self._file_info_widget)
        fi_row.setContentsMargins(8, 7, 8, 7)
        fi_row.setSpacing(7)

        pdf_icon = QLabel("📄")
        pdf_icon.setStyleSheet("color: #30D158; font-size: 14px; background: transparent; border: none;")
        fi_row.addWidget(pdf_icon)

        fi_text = QVBoxLayout()
        fi_text.setSpacing(1)
        self._file_info_label = QLabel("")
        self._file_info_label.setStyleSheet(
            "color: #E5E5EA; font-family: system-ui, -apple-system, sans-serif; font-size: 10px; font-weight: 500; background: transparent; border: none;"
        )
        self._detection_status_label = QLabel("")
        self._detection_status_label.setStyleSheet(
            "color: #48484A; font-family: system-ui, -apple-system, sans-serif; font-size: 9px; background: transparent; border: none;"
        )
        fi_text.addWidget(self._file_info_label)
        fi_text.addWidget(self._detection_status_label)
        fi_row.addLayout(fi_text, 1)

        self._btn_clear = PushButton("✕")
        self._btn_clear.setFixedSize(20, 20)
        self._btn_clear.setStyleSheet(
            "QPushButton { background: transparent; border: none; color: #636366; font-size: 11px; }"
            "QPushButton:hover { color: #FF453A; }"
        )
        self._btn_clear.clicked.connect(self._clear_bewerbung_pdf)
        fi_row.addWidget(self._btn_clear, 0, Qt.AlignVCenter)

        vbox.addWidget(self._file_info_widget)

        # ── Export single PDF ─────────────────────────────────────────────────
        self._btn_export_pdf = PushButton("Export PDF  →")
        self._btn_export_pdf.setFixedHeight(32)
        self._btn_export_pdf.setCursor(Qt.PointingHandCursor)
        self._btn_export_pdf.setStyleSheet("""
            QPushButton {
                background: rgba(48,209,88,0.10);
                border: 1px solid rgba(48,209,88,0.25);
                border-radius: 8px;
                color: #30D158;
                font-family: system-ui, -apple-system, sans-serif;
                font-size: 10px; font-weight: 500;
                letter-spacing: 0.04em;
            }
            QPushButton:hover { background: rgba(48,209,88,0.18); }
        """)
        self._btn_export_pdf.clicked.connect(self._action_export_bewerbungsmappe)
        vbox.addWidget(self._btn_export_pdf)

        # ── Batch export row ──────────────────────────────────────────────────
        self._btn_export_batch_pdf = PushButton("Batch export all visible  →")
        self._btn_export_batch_pdf.setFixedHeight(28)
        self._btn_export_batch_pdf.setCursor(Qt.PointingHandCursor)
        self._btn_export_batch_pdf.setStyleSheet("""
            QPushButton {
                background: transparent;
                border: 1px solid rgba(255,255,255,0.08);
                border-radius: 7px;
                color: #636366;
                font-family: system-ui, -apple-system, sans-serif;
                font-size: 9px; font-weight: 500;
                letter-spacing: 0.04em;
            }
            QPushButton:hover { border-color: rgba(255,255,255,0.18); color: #AEAEB2; }
        """)
        self._btn_export_batch_pdf.clicked.connect(self._action_export_bewerbungsmappe_batch)
        vbox.addWidget(self._btn_export_batch_pdf)

        layout.addWidget(card)


    def _on_pdf_dropped(self, path: str) -> None:
        self._load_bewerbung_pdf(Path(path))

    def _on_pdf_browse(self) -> None:
        start_dir = QStandardPaths.writableLocation(QStandardPaths.DocumentsLocation)
        if not start_dir:
            start_dir = str(Path.home())
        path, _ = QFileDialog.getOpenFileName(
            self,
            "Select Bewerbung PDF",
            start_dir,
            "PDF Files (*.pdf)"
        )
        if path:
            self._load_bewerbung_pdf(Path(path))

    def _open_bewerbung_pdf(self) -> None:
        if not self._bewerbung_pdf_path:
            self._show_error("No PDF loaded", "Add a Bewerbung PDF first.")
            return
        QDesktopServices.openUrl(QUrl.fromLocalFile(str(self._bewerbung_pdf_path)))

    def _run_bewerbung_export(self) -> None:
        self._set_mappe_generating(True)
        try:
            self._action_export_bewerbungsmappe()
        finally:
            self._set_mappe_generating(False)

    def _set_mappe_generating(self, generating: bool) -> None:
        if not hasattr(self, "_btn_bewerbung_batch"):
            return
        icon_name = "pause.svg" if generating else "play.svg"
        self._btn_bewerbung_batch.setIcon(_render_tinted_icon(icon_name, 14, "#FFFFFF"))

    def _load_bewerbung_pdf(self, path: Path) -> None:
        if not path.exists():
            return
        
        self._bewerbung_pdf_path = path

        if hasattr(self, "_file_info_widget"):
            self._file_info_widget.setVisible(True)
        if hasattr(self, "_drop_zone"):
            self._drop_zone.setVisible(False)
        if hasattr(self, "_file_info_label"):
            self._file_info_label.setText(path.name)
        if hasattr(self, "_detection_status_label"):
            self._detection_status_label.setText("Analyzing PDF...")

        if hasattr(self, "_btn_view_mappe"):
            self._btn_view_mappe.setEnabled(False)
        if hasattr(self, "_btn_generate"):
            self._btn_generate.setEnabled(False)

        self._pdf_worker = PdfDetectWorker(path, self)
        self._pdf_worker.detection_done.connect(self._on_pdf_detected)
        self._pdf_worker.start()

    def _on_pdf_detected(self, num_pages: int, detected_idx: int, was_detected: bool, error_msg: str) -> None:
        if error_msg:
            self._clear_bewerbung_pdf()
            self._show_error("Error loading PDF", error_msg)
            return

        # self._bewerbung_anschreiben_page is now managed by the UI input and user's config
        
        if was_detected:
            status_text = f"{num_pages} pages · letter p.{detected_idx + 1}"
        else:
            status_text = f"{num_pages} pages · default p.{detected_idx + 1}"

        if hasattr(self, "_detection_status_label"):
            self._detection_status_label.setText(status_text)

        if hasattr(self, "_bewerbung_status_label"):
            self._update_pdf_status_ui(True)
            self._bewerbung_status_label.setToolTip(str(self._bewerbung_pdf_path))

        if hasattr(self, "_btn_view_mappe"):
            self._btn_view_mappe.setEnabled(True)
        if hasattr(self, "_btn_generate"):
            self._btn_generate.setEnabled(True)

        if self._bewerbung_pdf_path:
            self._save_persisted_pdf_path(str(self._bewerbung_pdf_path))

    def _update_pdf_status_ui(self, loaded: bool, building: bool = False) -> None:
        if not hasattr(self, "_bewerbung_status_label"):
            return
        
        if building:
            self._bewerbung_status_label.setText("Building…")
            self._bewerbung_status_label.setStyleSheet("""
                QLabel {
                    background: #1A2A3A;
                    border: 0.5px solid #0A84FF;
                    border-radius: 6px;
                    color: #0A84FF;
                    padding: 3px 10px;
                    font-family: system-ui, -apple-system, "SF Pro Text", sans-serif;
                    font-size: 11px;
                    font-weight: 400;
                }
            """)
        elif loaded:
            self._bewerbung_status_label.setText("PDF ready")
            self._bewerbung_status_label.setStyleSheet("""
                QLabel {
                    background: #1A3A1A;
                    border: 0.5px solid #30A46C;
                    border-radius: 6px;
                    color: #30A46C;
                    padding: 3px 10px;
                    font-family: system-ui, -apple-system, "SF Pro Text", sans-serif;
                    font-size: 11px;
                    font-weight: 400;
                }
            """)
        else:
            self._bewerbung_status_label.setText("No PDF")
            self._bewerbung_status_label.setStyleSheet("""
                QLabel {
                    background: #2C2C2E;
                    border: 0.5px solid #3A3A3C;
                    border-radius: 6px;
                    color: #8E8E93;
                    padding: 3px 10px;
                    font-family: system-ui, -apple-system, "SF Pro Text", sans-serif;
                    font-size: 11px;
                    font-weight: 400;
                }
            """)

    def _clear_bewerbung_pdf(self) -> None:
        self._bewerbung_pdf_path = None
        # Don't reset page when clearing PDF, keep user's preference
        
        if hasattr(self, "_bewerbung_status_label"):
            self._update_pdf_status_ui(False)
            self._bewerbung_status_label.setToolTip("")
            
        if hasattr(self, "_btn_view_mappe"):
            self._btn_view_mappe.setEnabled(False)
        if hasattr(self, "_btn_generate"):
            self._btn_generate.setEnabled(False)

        self._save_persisted_pdf_path("")

    def _save_anschreiben_page(self):
        try:
            val = int(self._page_input.text())
            if val < 1: val = 1
            from ..core.config import config_manager
            config_manager.update(bewerbung_anschreiben_page=val)
            self._bewerbung_anschreiben_page = val - 1
        except ValueError:
            pass

    def _preview_merged_pdf(self):
        if not self._selected_record:
            self._show_error("Preview Failed", "No lead selected.")
            return
            
        import tempfile, os
        from PySide6.QtGui import QDesktopServices
        from PySide6.QtCore import QUrl
        
        self._status.setText("Generating preview...")
        QCoreApplication.processEvents()
        
        try:
            filled_text = self._editor.toPlainText()
            
            fd, tmp_path = tempfile.mkstemp(suffix=".pdf", prefix=f"preview_")
            os.close(fd)
            tmp_file = Path(tmp_path)
            
            if self._bewerbung_pdf_path:
                self._export_bewerbungsmappe(self._selected_record, filled_text, tmp_file)
            else:
                pdf_bytes = self._render_letter_as_pdf_page(filled_text)
                with open(tmp_file, "wb") as f:
                    f.write(pdf_bytes)
            
            QDesktopServices.openUrl(QUrl.fromLocalFile(str(tmp_file)))
            self._status.setText("Preview opened.")
        except Exception as e:
            self._show_error("Preview Error", str(e))

    def _render_letter_as_pdf_page(self, text: str) -> bytes:
        import io
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_LEFT, TA_JUSTIFY, TA_CENTER, TA_RIGHT
        from reportlab.platypus.flowables import KeepTogether

        buffer = io.BytesIO()
        
        # Modern Minimalist layout: tighter margins to ensure signature stays on page 1
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=10*mm,
            leftMargin=10*mm,
            topMargin=8*mm,
            bottomMargin=5*mm,
        )
        
        pdf_settings = self._load_pdf_settings()
        font_name = pdf_settings.get("font", "Helvetica")
        
        # Enforce highly readable defaults since UI is removed
        font_size = 11
        leading = 14
        alignment = TA_JUSTIFY
        
        styles = getSampleStyleSheet()
        style_normal = styles["Normal"]
        style_normal.fontName = font_name
        style_normal.fontSize = font_size
        style_normal.leading = leading
        style_normal.alignment = alignment
        
        style_h1 = ParagraphStyle(
            "H1",
            parent=style_normal,
            fontName=f"{font_name}-Bold" if font_name != "Times-Roman" else "Times-Bold",
            fontSize=30,
            leading=34,
            spaceAfter=4,
            alignment=TA_CENTER,
        )
        
        style_h2 = ParagraphStyle(
            "H2",
            parent=style_normal,
            fontName=f"{font_name}-Bold" if font_name != "Times-Roman" else "Times-Bold",
            fontSize=12,
            leading=14,
            spaceAfter=8,
            alignment=TA_CENTER,
        )
        
        style_right = ParagraphStyle(
            "Right",
            parent=style_normal,
            alignment=TA_RIGHT,
        )
        
        style_bold = ParagraphStyle(
            "Bold",
            parent=style_normal,
            fontName=f"{font_name}-Bold" if font_name != "Times-Roman" else "Times-Bold",
        )
        
        story = []
        lines = text.splitlines()
        
        if not lines:
            doc.build([Paragraph("Empty", style_normal)])
            return buffer.getvalue()
            
        # Parse the modern layout:
        idx = 0
        while idx < len(lines) and not lines[idx].strip():
            idx += 1
            
        if idx < len(lines):
            story.append(Paragraph(lines[idx].strip().upper(), style_h1))
            idx += 1
            
        while idx < len(lines) and not lines[idx].strip():
            idx += 1
            
        if idx < len(lines):
            story.append(Paragraph(lines[idx].strip(), style_h2))
            idx += 1
            
        # Lines 3, 4, etc. (normal paragraphs)
        is_first_paragraph = True
        for i in range(idx, len(lines)):
            line = lines[i].strip()
            if not line:
                story.append(Spacer(1, 10))
            else:
                if is_first_paragraph:
                    story.append(Paragraph(line, style_right))
                    is_first_paragraph = False
                elif "bewerbung um eine ausbildung" in line.lower() or "bewerbung als" in line.lower():
                    style_subj = ParagraphStyle("Subj", parent=style_bold, fontSize=13, spaceAfter=2)
                    story.append(Paragraph(line, style_subj))
                elif "freundlichen" in line.lower() and "gr" in line.lower():
                    story.append(Paragraph(line, style_normal))
                    story.append(Spacer(1, 2*mm)) # reduced from 5mm
                    
                    identity_block = []
                    
                    # Synchronous fallback to ensure we get the path even if background thread failed
                    sig_val = getattr(self, "_signature_image_path", "")
                    if not sig_val:
                        try:
                            import sqlite3
                            conn = sqlite3.connect(get_memory_db_path(), timeout=10.0)
                            row = conn.execute("SELECT value FROM settings WHERE key = 'signature_image_path'").fetchone()
                            if row and row[0]:
                                sig_val = row[0]
                            conn.close()
                        except Exception:
                            pass
                    
                    if sig_val:
                        sig_path = Path(sig_val)
                        if sig_path.exists():
                            try:
                                # Preprocess the image using Pillow to fix ReportLab transparency crash
                                from PIL import Image as PILImage
                                with PILImage.open(str(sig_path)) as pil_img:
                                    pil_img = pil_img.convert("RGBA")
                                    # Create white background to eliminate transparency
                                    bg = PILImage.new("RGBA", pil_img.size, (255, 255, 255, 255))
                                    out = PILImage.alpha_composite(bg, pil_img)
                                    rgb_img = out.convert("RGB")
                                    
                                    # Save to temporary buffer as JPEG
                                    img_buffer = io.BytesIO()
                                    rgb_img.save(img_buffer, format="JPEG")
                                    img_buffer.seek(0)
                                    
                                    img = Image(img_buffer, width=55*mm, height=20*mm, kind="proportional")
                                    img.hAlign = 'LEFT'
                                    identity_block.append(img)
                            except Exception as e:
                                identity_block.append(Paragraph(f"[Error processing signature image: {e}]", style_normal))
                        else:
                            identity_block.append(Paragraph(f"[Signature file not found: {sig_path}]", style_normal))
                    else:
                        identity_block.append(Paragraph("[No signature path saved in DB]", style_normal))
                                
                    sender_name = self._cached_sender_settings.get("name", "")
                    if not sender_name:
                        try:
                            import sqlite3
                            conn = sqlite3.connect(get_memory_db_path(), timeout=10.0)
                            row = conn.execute("SELECT value FROM settings WHERE key = 'sender_name'").fetchone()
                            if row and row[0]:
                                sender_name = row[0]
                            conn.close()
                        except Exception:
                            pass

                    if sender_name:
                        style_name = ParagraphStyle("Name", parent=style_bold, fontSize=13)
                        identity_block.append(Spacer(1, 1*mm))
                        identity_block.append(Paragraph(sender_name, style_name))
                    else:
                        identity_block.append(Paragraph("[No sender name set]", style_normal))
                        
                    story.append(KeepTogether(identity_block))
                else:
                    story.append(Paragraph(line, style_normal))

        doc.build(story)
        return buffer.getvalue()

    def _export_bewerbungsmappe(self, record: LeadRecord, letter_text: str, output_path: Path | None = None) -> Path:
        import pypdf
        import io

        if not self._bewerbung_pdf_path:
            raise ValueError("Upload your Bewerbung PDF first")

        letter_bytes = self._render_letter_as_pdf_page(letter_text)
        reader = pypdf.PdfReader(self._bewerbung_pdf_path)
        if reader.is_encrypted:
            raise ValueError("PDF is password-protected. Please unlock it first.")
            
        num_pages = len(reader.pages)
        letter_reader = pypdf.PdfReader(io.BytesIO(letter_bytes))
        
        writer = pypdf.PdfWriter()
        target_idx = self._bewerbung_anschreiben_page
        appended_at_end = False
        
        import sys
        old_limit = sys.getrecursionlimit()
        sys.setrecursionlimit(max(10000, old_limit))
        
        try:
            if target_idx < 0 or target_idx >= num_pages:
                appended_at_end = True
                for page in reader.pages:
                    writer.add_page(page)
                for new_page in letter_reader.pages:
                    writer.add_page(new_page)
            else:
                for i in range(num_pages):
                    if i == target_idx:
                        for new_page in letter_reader.pages:
                            writer.add_page(new_page)
                    else:
                        writer.add_page(reader.pages[i])
        finally:
            sys.setrecursionlimit(old_limit)
                    
        beruf = self._cached_sender_settings.get("beruf", "") or record.job_title or "Ausbildung"
        sender = self._load_sender_settings().get("name", "")
        if not sender:
            from ..core.config import config_manager
            sender = config_manager.settings.email_from_name
        if not sender:
            sender = "Bewerber"
        
        def sanitize(v):
            return re.sub(r'[<>:"/\\|?*]', '_', v)
            
        beruf_san = sanitize(beruf)
        sender_san = sanitize(sender)
        
        if output_path is None:
            filename = f"Bewerbung als {beruf_san} - {sender_san}.pdf"
            output_path = get_exports_dir() / filename

        
        with open(output_path, "wb") as f:
            writer.write(f)
            
        if appended_at_end:
            self._show_error(
                "Export Warning",
                f"Page {target_idx + 1} not found — Anschreiben appended at end."
            )
            
        return output_path

    def _export_bewerbungsmappe_batch(self, out_dir: Path | None = None, records: list = None) -> None:
        import pypdf
        import io

        if not self._bewerbung_pdf_path:
            raise ValueError("Upload your Bewerbung PDF first")

        reader = pypdf.PdfReader(self._bewerbung_pdf_path)
        if reader.is_encrypted:
            raise ValueError("PDF is password-protected. Please unlock it first.")
            
        num_pages = len(reader.pages)
        target_idx = self._bewerbung_anschreiben_page
        records = records if records is not None else self._pending_lead_records
        total = len(records)
        sender = self._load_sender_settings().get("name", "")
        if not sender:
            from ..core.config import config_manager
            sender = config_manager.settings.email_from_name
        if not sender:
            sender = "Bewerber"
        
        def sanitize(v):
            return re.sub(r'[<>:"/\\|?*]', '_', v)
            
        sender_san = sanitize(sender)
        warnings = []
        self._progress_info_bar = None
        
        try:
            for idx, record in enumerate(records):
                progress_text = f"Exporting {idx + 1} / {total}…"
                if self._progress_info_bar:
                    self._progress_info_bar.close()
                
                self._progress_info_bar = InfoBar.info(
                    title="Batch Export",
                    content=progress_text,
                    orient=Qt.Horizontal,
                    isClosable=False,
                    position=InfoBarPosition.TOP,
                    duration=-1,
                    parent=self
                )
                QCoreApplication.processEvents()
                
                state = self._states.get(record.id)
                letter_text = (state.letter_text if state else None) or self._assemble_letter(record)
                
                letter_bytes = self._render_letter_as_pdf_page(letter_text)
                letter_reader = pypdf.PdfReader(io.BytesIO(letter_bytes))
                new_page = letter_reader.pages[0]
                
                writer = pypdf.PdfWriter()
                appended_at_end = False
                
                if target_idx < 0 or target_idx >= num_pages:
                    appended_at_end = True
                    for page in reader.pages:
                        writer.add_page(page)
                    writer.add_page(new_page)
                else:
                    for i in range(num_pages):
                        if i == target_idx:
                            writer.add_page(new_page)
                        else:
                            writer.add_page(reader.pages[i])
                
                beruf = self._cached_sender_settings.get("beruf", "") or record.job_title or "Ausbildung"
                firma = record.company_name or "Firma"
                
                beruf_san = sanitize(beruf)
                firma_san = sanitize(firma)
                
                filename = f"Bewerbung als {beruf_san} - {sender_san} @ {firma_san}.pdf"
                if out_dir:
                    output_path = out_dir / filename
                else:
                    output_path = get_exports_dir() / filename
                
                with open(output_path, "wb") as f:
                    writer.write(f)
                    
                if appended_at_end:
                    warnings.append(record.company_name or "Unknown Company")
                    
            if self._progress_info_bar:
                self._progress_info_bar.close()
                self._progress_info_bar = None
                
            self._show_success("Batch Export Completed", f"Successfully exported {total} PDFs.")
            
            if warnings:
                self._show_error(
                    "Export Warning",
                    f"For {len(warnings)} leads, page {target_idx + 1} was out of range. Anschreiben was appended at the end."
                )
                
        except Exception as e:
            if self._progress_info_bar:
                self._progress_info_bar.close()
                self._progress_info_bar = None
            raise e

    def _action_export_bewerbungsmappe(self) -> None:
        if not self._selected_record:
            self._show_error("Export Failed", "Select a lead first")
            return
            
        if not self._bewerbung_pdf_path:
            self._show_error("Export Failed", "Upload your Bewerbung PDF first")
            return
            
        from ..core.security import LicenseManager
        if not LicenseManager.can_export_pdf(1):
            status = LicenseManager.get_pdf_trial_status()
            from qfluentwidgets import InfoBar, InfoBarPosition
            InfoBar.warning(
                title="Free Limit Reached",
                content=f"You have reached your free limit of {status['total']} PDFs per day. Please upgrade to Pro.",
                orient=Qt.Horizontal,
                isClosable=True,
                position=InfoBarPosition.TOP,
                duration=6000,
                parent=self.window()
            )
            return
            
        try:
            import pypdf
        except ImportError:
            self._show_error("Dependency Missing", "Install pypdf:  pip install pypdf")
            return
            
        try:
            import reportlab
        except ImportError:
            self._show_error("Dependency Missing", "Install reportlab:  pip install reportlab")
            return

        try:
            beruf = self._cached_sender_settings.get("beruf", "") or self._selected_record.job_title or "Ausbildung"
            sender = self._load_sender_settings().get("name", "Bewerber")
            beruf_san = re.sub(r'[<>:"/\\|?*]', '_', beruf)
            sender_san = re.sub(r'[<>:"/\\|?*]', '_', sender)
            default_name = f"Bewerbung als {beruf_san} - {sender_san}.pdf"
            
            out_file, _ = QFileDialog.getSaveFileName(
                self,
                "Save Bewerbungsmappe PDF",
                default_name,
                "PDF Files (*.pdf)"
            )
            
            if not out_file:
                return
                
            letter_text = self._editor.toPlainText()
            out_path = self._export_bewerbungsmappe(self._selected_record, letter_text, output_path=Path(out_file))
            LicenseManager.record_pdf_export(1)
            self._show_success("Exported", out_path.name)
        except Exception as e:
            self._show_error("Export Failed", str(e))

    def _action_export_bewerbungsmappe_batch(self) -> None:
        if not self._pending_lead_records:
            self._show_error("Nothing to export", "No leads match current filter.")
            return
            
        if not self._bewerbung_pdf_path:
            self._show_error("Export Failed", "Upload your Bewerbung PDF first")
            return
            
        from ..core.security import LicenseManager
        count = len(self._pending_lead_records)
        status = LicenseManager.get_pdf_trial_status()
        export_records = self._pending_lead_records
        show_limit_warning = False

        if not LicenseManager.is_active() and count > status['remaining']:
            if status['remaining'] <= 0:
                from qfluentwidgets import InfoBar, InfoBarPosition
                InfoBar.warning(
                    title="Free Limit Reached",
                    content=f"Batch export of {count} PDFs exceeds your daily limit of 0. Please upgrade to Pro.",
                    orient=Qt.Horizontal,
                    isClosable=True,
                    position=InfoBarPosition.TOP,
                    duration=6000,
                    parent=self.window()
                )
                return
            else:
                export_records = self._pending_lead_records[:status['remaining']]
                show_limit_warning = True
            
        try:
            import pypdf
        except ImportError:
            self._show_error("Dependency Missing", "Install pypdf:  pip install pypdf")
            return
            
        try:
            import reportlab
        except ImportError:
            self._show_error("Dependency Missing", "Install reportlab:  pip install reportlab")
            return
            
        try:
            out_dir = QFileDialog.getExistingDirectory(
                self,
                "Select Export Folder for Batch PDF",
                ""
            )
            
            if not out_dir:
                return
                
            self._export_bewerbungsmappe_batch(out_dir=Path(out_dir), records=export_records)
            LicenseManager.record_pdf_export(len(export_records))
            
            if show_limit_warning:
                from qfluentwidgets import InfoBar, InfoBarPosition
                InfoBar.warning(
                    title="Free Limit Reached",
                    content=f"Batch export of {count} leads exceeds your limit. Only {max(0, status['remaining'])} PDFs were generated. Please upgrade to Pro for unlimited usage.",
                    orient=Qt.Horizontal,
                    isClosable=True,
                    position=InfoBarPosition.TOP,
                    duration=6000,
                    parent=self.window()
                )
        except Exception as e:
            self._show_error("Batch Export Failed", str(e))

    def _action_export_and_send_batch(self) -> None:
        if not self._pending_lead_records:
            self._show_error("Export Failed", "No leads match current filter.")
            return
            
        if not self._bewerbung_pdf_path:
            self._show_error("Export Failed", "Upload your Bewerbung PDF first")
            return
            
        from ..core.security import LicenseManager
        count = len(self._pending_lead_records)
        status = LicenseManager.get_pdf_trial_status()
        export_records = self._pending_lead_records
        show_limit_warning = False

        if not LicenseManager.is_active() and count > status['remaining']:
            export_records = self._pending_lead_records[:max(0, status['remaining'])]
            show_limit_warning = True
            
        try:
            import pypdf
        except ImportError:
            self._show_error("Dependency Missing", "Install pypdf:  pip install pypdf")
            return
            
        try:
            import reportlab
        except ImportError:
            self._show_error("Dependency Missing", "Install reportlab:  pip install reportlab")
            return
            
        try:
            # We don't ask for a directory, we just save to the standard exports folder
            from ..core.config import get_exports_dir
            import shutil
            
            out_dir = get_exports_dir()
            
            # Copy the raw PDF to act as a fallback for leads beyond the trial limit
            raw_pdf_path = out_dir / "Bewerbung_Raw_Uploaded.pdf"
            try:
                if self._bewerbung_pdf_path and self._bewerbung_pdf_path.exists():
                    shutil.copy2(str(self._bewerbung_pdf_path), str(raw_pdf_path))
            except Exception as e:
                pass # Non-fatal if we can't copy it
                
            self._export_bewerbungsmappe_batch(out_dir=Path(out_dir), records=export_records)
            LicenseManager.record_pdf_export(len(export_records))
            
            # Emit ALL valid emails, even those that didn't get a custom PDF
            all_emails = [r.email for r in self._pending_lead_records if r.email and r.email.strip()]
            if all_emails:
                self.send_emails_to_sender.emit(all_emails)
            else:
                self._show_error("No Emails", "No valid emails found in the visible leads.")
                
            if show_limit_warning:
                from qfluentwidgets import InfoBar, InfoBarPosition
                InfoBar.warning(
                    title="Free Limit Reached",
                    content=f"Batch export of {count} leads exceeds your limit. Only {max(0, status['remaining'])} custom PDFs were generated. The rest will use your raw uploaded PDF as a fallback. Please upgrade to Pro for unlimited customization.",
                    orient=Qt.Horizontal,
                    isClosable=True,
                    position=InfoBarPosition.TOP,
                    duration=6000,
                    parent=self.window()
                )
        except Exception as e:
            self._show_error("Batch Export & Send Failed", str(e))


class TemplateEditorDialog(QDialog):
    def __init__(self, parent, full_template_text: str):
        super().__init__(parent)
        self.parent_page = parent
        self.setWindowTitle("Edit Anschreiben Template")
        self.setMinimumSize(850, 650)
        self.setStyleSheet("QDialog { background: #1C1C1E; }")
        
        self._full_text = full_template_text
        self._header_part = ""
        self._is_body_mode = False
        
        layout = QVBoxLayout(self)
        layout.setContentsMargins(18, 18, 18, 18)
        layout.setSpacing(16)
        
        from qfluentwidgets import SegmentedWidget, PushButton, PrimaryPushButton, InfoBar, InfoBarPosition
        
        # Apple-style Segmented Control Container
        segment_container = QFrame()
        segment_container.setStyleSheet("""
            QFrame {
                background: rgba(255, 255, 255, 0.04);
                border-radius: 8px;
                padding: 3px;
            }
        """)
        segment_layout = QHBoxLayout(segment_container)
        segment_layout.setContentsMargins(3, 3, 3, 3)
        segment_layout.setSpacing(0)
        
        from PySide6.QtWidgets import QPushButton
        self.btn_body = QPushButton("Edit Body Only")
        self.btn_full = QPushButton("Edit Full Layout")
        self.btn_body.setCursor(Qt.CursorShape.PointingHandCursor)
        self.btn_full.setCursor(Qt.CursorShape.PointingHandCursor)
        
        # We define a helper to update styles
        def _update_segment_styles(is_body: bool):
            active_style = """
                QPushButton {
                    background: rgba(255, 255, 255, 0.12);
                    border: 1px solid rgba(255, 255, 255, 0.05);
                    border-radius: 6px;
                    color: #FFFFFF;
                    font-family: system-ui, -apple-system, sans-serif;
                    font-size: 11px;
                    font-weight: 600;
                    letter-spacing: 1px;
                    padding: 8px 16px;
                    min-height: 28px;
                }
                QPushButton:hover { background: rgba(255, 255, 255, 0.15); }
            """
            inactive_style = """
                QPushButton {
                    background: transparent;
                    border: none;
                    color: #8E8E93;
                    font-family: system-ui, -apple-system, sans-serif;
                    font-size: 11px;
                    font-weight: 600;
                    letter-spacing: 1px;
                    padding: 8px 16px;
                    min-height: 28px;
                }
                QPushButton:hover { color: #FFFFFF; background: rgba(255, 255, 255, 0.05); border-radius: 6px; }
            """
            self.btn_body.setStyleSheet(active_style if is_body else inactive_style)
            self.btn_full.setStyleSheet(inactive_style if is_body else active_style)
            
        self._update_segment_styles = _update_segment_styles
        self.btn_body.clicked.connect(lambda: self._on_segment_changed("body"))
        self.btn_full.clicked.connect(lambda: self._on_segment_changed("full"))
        
        segment_layout.addWidget(self.btn_body)
        segment_layout.addWidget(self.btn_full)
        
        # Center the segmented control
        seg_wrapper = QHBoxLayout()
        seg_wrapper.addStretch(1)
        seg_wrapper.addWidget(segment_container)
        seg_wrapper.addStretch(1)
        layout.addLayout(seg_wrapper)
        
        split = QHBoxLayout()
        split.setSpacing(14)
        
        # Editor Side
        editor_col = QVBoxLayout()
        editor_col.setSpacing(8)
        
        self.hint = QLabel("Write your Anschreiben below.")
        self.hint.setWordWrap(True)
        self.hint.setStyleSheet(
            "color: #AEAEB2; font-family: system-ui, -apple-system, sans-serif; "
            "font-size: 12px; background: transparent;"
        )
        editor_col.addWidget(self.hint)
        
        self.editor = QTextEdit()
        # Enhanced text editor style to match app aesthetic with elegant framing
        self.editor.setStyleSheet("""
            QTextEdit {
                background: rgba(255, 255, 255, 0.03);
                border: 1px solid rgba(255, 255, 255, 0.08);
                border-radius: 8px;
                color: #E5E5EA;
                font-family: system-ui, -apple-system, sans-serif;
                font-size: 14px;
                padding: 16px;
            }
            QTextEdit:focus {
                border: 1px solid rgba(255, 255, 255, 0.15);
                background: rgba(255, 255, 255, 0.05);
            }
        """)
        
        # Adjust line height using QTextBlockFormat for a cleaner look
        from PySide6.QtGui import QTextBlockFormat
        cursor = self.editor.textCursor()
        cursor.select(cursor.SelectionType.Document)
        block_fmt = QTextBlockFormat()
        # 4 corresponds to ProportionalHeight in Qt
        block_fmt.setLineHeight(150.0, 4)
        cursor.mergeBlockFormat(block_fmt)
        self.editor.setTextCursor(cursor)
        
        editor_col.addWidget(self.editor, 1)
        split.addLayout(editor_col, 2)
        
        # Placeholders Side (Wrapped in a QWidget to hide/show)
        self.ref_widget = QWidget()
        ref_col = QVBoxLayout(self.ref_widget)
        ref_col.setContentsMargins(0, 0, 0, 0)
        ref_col.setSpacing(6)
        
        ref_title = QLabel("PLACEHOLDERS")
        ref_title.setStyleSheet(
            "color: #8E8E93; font-family: system-ui, -apple-system, sans-serif; "
            "font-size: 10px; font-weight: 500; background: transparent;"
        )
        ref_col.addWidget(ref_title)
        
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setStyleSheet("QScrollArea { border: none; background: transparent; }")
        scroll_content = QWidget()
        scroll_content.setStyleSheet("background: transparent;")
        scroll_layout = QVBoxLayout(scroll_content)
        scroll_layout.setContentsMargins(0, 0, 4, 0)
        scroll_layout.setSpacing(8)
        
        for ph, desc in PLACEHOLDER_REFERENCE:
            card = PushButton(self)
            card.setText("")
            
            card_layout = QVBoxLayout(card)
            card_layout.setContentsMargins(12, 10, 12, 10)
            card_layout.setSpacing(2)
            
            lbl_token = QLabel(ph)
            lbl_token.setStyleSheet("color: #0A84FF; font-family: 'Menlo', monospace; font-size: 12px; font-weight: bold; background: transparent; border: none;")
            
            lbl_desc = QLabel(desc)
            lbl_desc.setWordWrap(True)
            lbl_desc.setStyleSheet("color: #98989D; font-family: system-ui, -apple-system, sans-serif; font-size: 11px; background: transparent; border: none;")
            
            card_layout.addWidget(lbl_token)
            card_layout.addWidget(lbl_desc)
            
            card.setStyleSheet("""
                PushButton {
                    background-color: rgba(255,255,255,0.03); 
                    border: 1px solid rgba(255,255,255,0.08); 
                    border-radius: 6px;
                    text-align: left;
                }
                PushButton:hover { 
                    background-color: rgba(255,255,255,0.06); 
                    border: 1px solid rgba(255,255,255,0.15); 
                }
            """)
            card.setCursor(Qt.CursorShape.PointingHandCursor)
            card.clicked.connect(lambda checked, p=ph: self.editor.textCursor().insertText(p))
            scroll_layout.addWidget(card)
            
        scroll_layout.addStretch(1)
        scroll.setWidget(scroll_content)
        ref_col.addWidget(scroll)
        
        split.addWidget(self.ref_widget, 1)
        layout.addLayout(split, 1)
        
        # Bottom Actions
        buttons = QHBoxLayout()
        buttons.addStretch(1)
        
        from PySide6.QtWidgets import QPushButton
        cancel = QPushButton("Cancel")
        save = QPushButton("Save Template")
        
        cancel.setCursor(Qt.CursorShape.PointingHandCursor)
        cancel.setFixedHeight(32)
        cancel.setStyleSheet("""
            QPushButton {
                background: rgba(255, 255, 255, 0.1);
                border: none;
                border-radius: 6px;
                color: white;
                font-family: 'SF Pro Text', 'PT Root UI', sans-serif;
                font-size: 13px;
                font-weight: 500;
                padding: 0 16px;
            }
            QPushButton:hover { background: rgba(255, 255, 255, 0.15); }
        """)
        
        save.setCursor(Qt.CursorShape.PointingHandCursor)
        save.setFixedHeight(32)
        save.setStyleSheet("""
            QPushButton {
                background: #0A84FF;
                border: none;
                border-radius: 6px;
                color: #FFFFFF;
                font-family: 'SF Pro Text', 'PT Root UI', sans-serif;
                font-size: 13px;
                font-weight: 600;
                padding: 0 16px;
            }
            QPushButton:hover { background: #007AFF; }
        """)
            
        cancel.clicked.connect(self.reject)
        save.clicked.connect(self.accept)
        buttons.addWidget(cancel)
        buttons.addWidget(save)
        layout.addLayout(buttons)
        
        # Initialize
        self._on_segment_changed("body")
        
    def _on_segment_changed(self, item_key):
        import re
        if item_key == "body":
            if not self._is_body_mode and self.editor.toPlainText():
                self._full_text = self.editor.toPlainText()
                
            match = re.search(r"(?i)(Sehr geehrte[r]?.*?|Guten Tag.*?|\{\{ANREDE\}\}),?\s*\n+", self._full_text)
            if match:
                split_idx = match.end()
                self._header_part = self._full_text[:split_idx]
                body_part = self._full_text[split_idx:]
                self.editor.setPlainText(body_part.strip() + "\n")
                self._is_body_mode = True
                self.hint.setText("Write your Anschreiben below.")
                self.ref_widget.hide()
                self._update_segment_styles(True)
            else:
                self._on_segment_changed("full")
                from qfluentwidgets import InfoBar, InfoBarPosition
                InfoBar.warning(
                    title="Layout Not Standard",
                    content="Could not detect standard greeting. Reverting to Full Layout.",
                    parent=self.parent_page,
                    position=InfoBarPosition.TOP,
                    duration=3000
                )
        else:
            if self._is_body_mode:
                self._full_text = self._header_part + "\n" + self.editor.toPlainText().strip() + "\n"
            self.editor.setPlainText(self._full_text)
            self._is_body_mode = False
            self.hint.setText("Edit Full Layout. Placeholders are replaced automatically.")
            self.ref_widget.show()
            self._update_segment_styles(False)
            
    def get_template_text(self):
        if self._is_body_mode:
            return self._header_part + "\n" + self.editor.toPlainText().strip() + "\n"
        return self.editor.toPlainText().strip() + "\n"
