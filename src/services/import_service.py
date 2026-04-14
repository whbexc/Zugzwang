"""
ZUGZWANG - Import Service
Handles importing of external lists (CSV, XLSX, DOCX, TXT) and clipboard paste.
"""

import csv
import io
import re
import uuid
import openpyxl
from urllib.parse import urlparse
from collections import defaultdict
from enum import Enum
from typing import Optional, Any

from ..core.models import LeadRecord, SourceType
from ..core.logger import get_logger

logger = get_logger(__name__)

class ImportService:
    def __init__(self):
        self._url_regex = re.compile(
            r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+'
        )
        self._email_regex = re.compile(
            r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+'
        )

    def import_from_file(self, file_path: str) -> list[LeadRecord]:
        """Import leads from a supported file format."""
        logger.info(f"Importing leads from file: {file_path}")
        ext = file_path.lower().split('.')[-1]
        
        if ext == 'csv':
            return self._parse_csv(file_path)
        elif ext == 'xlsx':
            return self._parse_xlsx(file_path)
        elif ext == 'txt':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            return self._parse_text(content)
        elif ext == 'docx':
            return self._parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported import format: .{ext}")

    def import_from_text(self, text: str) -> list[LeadRecord]:
        """Import leads from raw text (e.g., clipboard)."""
        logger.info(f"Importing leads from clipboard/text ({len(text)} chars)")
        return self._parse_text(text)

    def _parse_csv(self, file_path: str) -> list[LeadRecord]:
        records = []
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            reader = csv.DictReader(f)
            # Find closest matching columns
            if not reader.fieldnames:
                return records
            
            headers = [h.lower() for h in reader.fieldnames]
            col_map = self._map_columns(headers)
            
            for row in reader:
                record_data = {}
                for key, mapped_field in col_map.items():
                    val = row.get(next(h for h in reader.fieldnames if h.lower() == key), "")
                    record_data[mapped_field] = val.strip() if val else None
                
                if any(record_data.values()): # Only if not entirely empty
                    records.append(self._build_record_from_dict(record_data))
                    
        return self._dedupe_records(records)
        
    def _parse_xlsx(self, file_path: str) -> list[LeadRecord]:
        records = []
        try:
            wb = openpyxl.load_workbook(file_path, data_only=True)
            sheet = wb.active
            rows = list(sheet.iter_rows(values_only=True))
            if len(rows) < 2:
                return records
                
            headers = [str(h).lower() if h else f"col_{i}" for i, h in enumerate(rows[0])]
            col_map = self._map_columns(headers)
            
            for row in rows[1:]:
                record_data = {}
                for i, cell_val in enumerate(row):
                    if i < len(headers) and headers[i] in col_map:
                        mapped_field = col_map[headers[i]]
                        record_data[mapped_field] = str(cell_val).strip() if cell_val is not None else None
                        
                if any(record_data.values()):
                    records.append(self._build_record_from_dict(record_data))
                    
            return self._dedupe_records(records)
        except Exception as e:
            logger.error(f"XLSX Import failed: {e}")
            raise ValueError(f"Could not read XLSX file: {e}")

    def _parse_docx(self, file_path: str) -> list[LeadRecord]:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required. Please install it (pip install python-docx) or import as txt instead.")
            
        doc = Document(file_path)
        content = "\n".join([p.text for p in doc.paragraphs])
        # Also grab tables just in case
        for table in doc.tables:
            for row in table.rows:
                content += "\n" + " ".join([cell.text for cell in row.cells])
                
        return self._parse_text(content)

    def _parse_text(self, text: str) -> list[LeadRecord]:
        """Extract emails and optional websites from raw text block."""
        records = []
        
        # We find all emails in the text block
        found_emails = list(set(self._email_regex.findall(text)))
        found_urls = list(set(self._url_regex.findall(text)))
        
        for email in found_emails:
            domain = email.split('@')[1] if '@' in email else None
            matched_url = None
            if domain:
                # Try to pair with a matching URL if we found one
                for url in found_urls:
                    if domain.lower() in url.lower():
                        matched_url = url
                        break
            
            record = LeadRecord(
                id=str(uuid.uuid4()),
                source_type=SourceType.MANUAL,
                email=email,
                website=matched_url,
                company_name=domain.capitalize().split('.')[0] if domain else "Imported Lead"
            )
            records.append(record)
            
        # Add remaining unmatched URLs as separate records
        matched_urls_set = {r.website for r in records if r.website}
        for url in found_urls:
            if url not in matched_urls_set:
                records.append(LeadRecord(
                    id=str(uuid.uuid4()),
                    source_type=SourceType.MANUAL,
                    website=url,
                    company_name="Imported Lead"
                ))
                
        return records

    def _map_columns(self, headers: list[str]) -> dict[str, str]:
        """Map generic user header names to LeadRecord attributes."""
        mapping = {}
        header_patterns = {
            "company_name": ["company", "name", "firma", "business", "unternehmen", "organization"],
            "email": ["email", "e-mail", "mail", "contact_email"],
            "phone": ["phone", "tel", "telefon", "mobile", "cell"],
            "website": ["website", "web", "url", "site", "domain"],
            "address": ["address", "adresse", "street", "straße", "str"],
            "city": ["city", "stadt", "ort", "location"],
            "postal_code": ["zip", "postal", "plz", "postcode"],
            "contact_person": ["contact", "person", "ansprechpartner", "name"],
            "linkedin": ["linkedin", "li"],
            "job_title": ["job", "title", "position"]
        }
        
        for header in headers:
            for field, patterns in header_patterns.items():
                if any(p in header for p in patterns) and field not in mapping.values():
                    mapping[header] = field
                    break
        return mapping

    def _build_record_from_dict(self, data: dict[str, Any]) -> LeadRecord:
        record = LeadRecord(id=str(uuid.uuid4()), source_type=SourceType.MANUAL)
        for k, v in data.items():
            if hasattr(record, k) and v:
                setattr(record, k, v)
                
        if not record.company_name:
            if record.website:
                # Try to extract company name from domain
                domain = urlparse(record.website).netloc or record.website
                domain = domain.replace("www.", "")
                record.company_name = domain.split('.')[0].capitalize()
            elif record.email:
                domain = record.email.split('@')[1] if '@' in record.email else ""
                record.company_name = domain.split('.')[0].capitalize() if domain else "Imported Lead"
            else:
                record.company_name = "Imported Lead"
                
        return record
        
    def _dedupe_records(self, records: list[LeadRecord]) -> list[LeadRecord]:
        seen = set()
        deduped = []
        for r in records:
            # We group by company and email or website
            key = f"{str(r.company_name).lower()}|{str(r.email).lower() if r.email else ''}|{str(r.website).lower() if r.website else ''}"
            if key not in seen:
                seen.add(key)
                deduped.append(r)
        return deduped
