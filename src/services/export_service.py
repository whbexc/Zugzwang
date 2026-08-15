"""
ZUGZWANG - Export Service
Handles all data export formats: CSV, Excel (XLSX), JSON, SQLite.
Supports filtered/full export and project save/load.
"""

from __future__ import annotations
import csv
from dataclasses import replace
import json
import sqlite3
import os
from datetime import datetime
from pathlib import Path
from typing import Optional

from ..core.events import event_bus
from ..core.logger import get_logger
from ..core.models import LeadRecord, ScrapingJob, SourceType

logger = get_logger(__name__)

# Column headers in display order
EXPORT_COLUMNS = [
    "id", "source_type", "company_name", "job_title", "category",
    "email", "email_source_page", "phone", "website", "contact_person",
    "address", "city", "region", "postal_code", "country",
    "rating", "review_count", "description", "publication_date",
    "source_url", "maps_url", "search_query", "scraped_at", "notes",
]


class ExportService:
    """
    Handles all export and persistence operations.
    All methods are synchronous (called from export worker thread).
    """

    def export_csv(self, records: list[LeadRecord], path: str) -> int:
        """Export records to CSV. Returns number of rows written."""
        logger.info(f"Exporting {len(records)} records to CSV: {path}")
        try:
            with open(path, "w", newline="", encoding="utf-8-sig") as f:
                writer = csv.DictWriter(
                    f,
                    fieldnames=EXPORT_COLUMNS,
                    extrasaction="ignore",
                )
                writer.writeheader()
                for record in records:
                    row = record.to_dict()
                    writer.writerow({col: row.get(col, "") for col in EXPORT_COLUMNS})
            event_bus.emit(event_bus.EXPORT_COMPLETED, format="csv", path=path, count=len(records))
            return len(records)
        except Exception as e:
            logger.error(f"CSV export failed: {e}")
            event_bus.emit(event_bus.EXPORT_FAILED, format="csv", error=str(e))
            raise

    def export_json(self, records: list[LeadRecord], path: str) -> int:
        """Export records to JSON."""
        logger.info(f"Exporting {len(records)} records to JSON: {path}")
        try:
            data = [record.to_dict() for record in records]
            with open(path, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=2, ensure_ascii=False, default=str)
            event_bus.emit(event_bus.EXPORT_COMPLETED, format="json", path=path, count=len(records))
            return len(records)
        except Exception as e:
            logger.error(f"JSON export failed: {e}")
            event_bus.emit(event_bus.EXPORT_FAILED, format="json", error=str(e))
            raise

    def export_excel(self, records: list[LeadRecord], path: str) -> int:
        """Export records to Excel XLSX format."""
        try:
            import openpyxl
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
            from openpyxl.utils import get_column_letter
        except ImportError:
            raise ImportError("openpyxl required for Excel export: pip install openpyxl")

        logger.info(f"Exporting {len(records)} records to Excel: {path}")
        try:
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Leads"

            # Header styling
            header_fill = PatternFill("solid", fgColor="1A2B3C")
            header_font = Font(bold=True, color="FFFFFF", name="Calibri", size=10)
            alt_fill = PatternFill("solid", fgColor="F0F4F8")
            thin = Side(border_style="thin", color="D0D8E0")
            border = Border(left=thin, right=thin, top=thin, bottom=thin)

            # Write headers
            for col_idx, col_name in enumerate(EXPORT_COLUMNS, 1):
                cell = ws.cell(row=1, column=col_idx, value=col_name.replace("_", " ").title())
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal="center", vertical="center")
                cell.border = border

            ws.row_dimensions[1].height = 20

            # Write data rows
            for row_idx, record in enumerate(records, 2):
                row_data = record.to_dict()
                fill = alt_fill if row_idx % 2 == 0 else None
                for col_idx, col_name in enumerate(EXPORT_COLUMNS, 1):
                    val = row_data.get(col_name, "")
                    if val is None:
                        val = ""
                    cell = ws.cell(row=row_idx, column=col_idx, value=str(val))
                    cell.font = Font(name="Calibri", size=9)
                    cell.border = border
                    if fill:
                        cell.fill = fill
                    # Hyperlink for URLs
                    if col_name in ("website", "source_url", "maps_url", "email_source_page") and val:
                        url = val if val.startswith("http") else f"mailto:{val}" if "@" in val else val
                        try:
                            cell.hyperlink = url
                            cell.font = Font(name="Calibri", size=9, color="0066CC", underline="single")
                        except Exception:
                            pass

            # Auto-size columns
            col_widths = {
                "company_name": 30, "email": 35, "website": 40, "address": 35,
                "job_title": 30, "phone": 18, "city": 18, "description": 50,
            }
            for col_idx, col_name in enumerate(EXPORT_COLUMNS, 1):
                width = col_widths.get(col_name, 15)
                ws.column_dimensions[get_column_letter(col_idx)].width = width

            # Freeze header row
            ws.freeze_panes = "A2"

            # Auto-filter
            ws.auto_filter.ref = ws.dimensions

            wb.save(path)
            event_bus.emit(event_bus.EXPORT_COMPLETED, format="xlsx", path=path, count=len(records))
            return len(records)
        except Exception as e:
            logger.error(f"Excel export failed: {e}")
            event_bus.emit(event_bus.EXPORT_FAILED, format="xlsx", error=str(e))
            raise

    def export_docx(self, records: list[LeadRecord], path: str) -> int:
        """Export records to a Word DOCX document."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx required for Word export: pip install python-docx")

        logger.info(f"Exporting {len(records)} records to Word: {path}")
        try:
            doc = Document()
            doc.add_heading("ZUGZWANG Export", level=1)
            doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"Records: {len(records)}")

            for index, record in enumerate(records, 1):
                doc.add_heading(f"{index}. {record.company_name or record.job_title or 'Lead'}", level=2)
                table = doc.add_table(rows=0, cols=2)
                table.style = "Table Grid"

                row_data = record.to_dict()
                for col in EXPORT_COLUMNS:
                    value = row_data.get(col, "")
                    if value in (None, ""):
                        continue
                    cells = table.add_row().cells
                    cells[0].text = col.replace("_", " ").title()
                    cells[1].text = str(value)

            doc.save(path)
            event_bus.emit(event_bus.EXPORT_COMPLETED, format="docx", path=path, count=len(records))
            return len(records)
        except Exception as e:
            logger.error(f"Word export failed: {e}")
            event_bus.emit(event_bus.EXPORT_FAILED, format="docx", error=str(e))
            raise

    def export_txt(self, records: list[LeadRecord], path: str) -> int:
        """Export emails from records to a plain text file, one per line."""
        logger.info(f"Exporting emails from {len(records)} records to TXT: {path}")
        try:
            emails = [r.email for r in records if r.email]
            with open(path, "w", encoding="utf-8") as f:
                if emails:
                    f.write("\n".join(emails) + "\n")
                else:
                    # If no emails, write a summary instead of an empty file
                    for index, record in enumerate(records, 1):
                        f.write(f"{index}. {record.company_name or 'Lead'}\n")
                        if record.website: f.write(f"   Website: {record.website}\n")
                        if record.phone: f.write(f"   Phone: {record.phone}\n")
                        f.write("\n")
            
            event_bus.emit(event_bus.EXPORT_COMPLETED, format="txt", path=path, count=len(records))
            return len(records)
        except Exception as e:
            logger.error(f"TXT export failed: {e}")
            event_bus.emit(event_bus.EXPORT_FAILED, format="txt", error=str(e))
            raise

    def save_project(self, job: ScrapingJob, path: str) -> None:
        """Save a scraping job (with results) to SQLite project file."""
        logger.info(f"Saving project to: {path}")
        # Increase timeout to 30s to avoid 'database is locked' during concurrent access
        conn = sqlite3.connect(path, timeout=30.0)
        try:
            conn.execute("PRAGMA journal_mode=WAL")
            conn.execute("PRAGMA synchronous=NORMAL") # Faster and safe enough with WAL
            self._init_db(conn)
            conn.execute(
                "INSERT OR REPLACE INTO jobs (id, config_json, status, stats_json, created_at, started_at, completed_at) VALUES (?,?,?,?,?,?,?)",
                (
                    job.id,
                    json.dumps(self._serialize_config(job)),
                    job.status.value,
                    json.dumps({
                        "total_found": job.total_found,
                        "total_emails": job.total_emails,
                        "total_websites": job.total_websites,
                        "total_errors": job.total_errors,
                    }),
                    job.created_at,
                    job.started_at,
                    job.completed_at,
                ),
            )
            for record in job.results:
                prepared = self._prepare_record(record)
                row = prepared.to_dict()
                conn.execute(
                    f"INSERT OR REPLACE INTO leads ({','.join(EXPORT_COLUMNS)}) VALUES ({','.join(['?']*len(EXPORT_COLUMNS))})",
                    [row.get(c, "") for c in EXPORT_COLUMNS],
                )
            conn.commit()
            logger.info(f"Project saved: {len(job.results)} leads")
        finally:
            conn.close()

    def load_project(self, path: str) -> tuple[Optional[dict], list[LeadRecord]]:
        """Load a full project (metadata + all leads). Returns (job_meta, records)."""
        job_meta = self.load_job_metadata(path)
        
        conn = sqlite3.connect(path)
        try:
            conn.execute("PRAGMA journal_mode=WAL")
            self._init_db(conn)
            leads_rows = conn.execute(f"SELECT {','.join(EXPORT_COLUMNS)} FROM leads").fetchall()
            records = []
            seen_ids: set[str] = set()
            for row in leads_rows:
                d = dict(zip(EXPORT_COLUMNS, row))
                try:
                    record = LeadRecord.from_dict(d).normalize()
                    record.id = record.stable_id()
                    if record.id in seen_ids:
                        continue
                    seen_ids.add(record.id)
                    records.append(record)
                except Exception as e:
                    logger.warning(f"Could not parse lead record: {e}")

            return job_meta, records
        finally:
            conn.close()

    def load_job_metadata(self, path: str) -> Optional[dict]:
        """Lightweight load of just the job metadata. Fast for dashboard list."""
        if not os.path.exists(path):
            return None
        logger.debug(f"Loading job metadata from: {path}")
        conn = sqlite3.connect(path)
        try:
            conn.execute("PRAGMA journal_mode=WAL")
            self._init_db(conn)
            job_row = conn.execute("SELECT * FROM jobs ORDER BY created_at DESC LIMIT 1").fetchone()
            if not job_row:
                return None
            cols = [d[0] for d in conn.execute("SELECT * FROM jobs LIMIT 0").description]
            return dict(zip(cols, job_row))
        except Exception as e:
            logger.warning(f"Failed to load job metadata from {path}: {e}")
            return None
        finally:
            conn.close()

    def _init_db(self, conn: sqlite3.Connection) -> None:
        conn.execute("""
            CREATE TABLE IF NOT EXISTS jobs (
                id TEXT PRIMARY KEY,
                config_json TEXT,
                status TEXT,
                stats_json TEXT,
                created_at TEXT,
                started_at TEXT,
                completed_at TEXT
            )
        """)
        cols_def = ", ".join(f"{c} TEXT" for c in EXPORT_COLUMNS)
        conn.execute(f"""
            CREATE TABLE IF NOT EXISTS leads (
                {cols_def},
                PRIMARY KEY (id)
            )
        """)
        existing = {
            row[1]
            for row in conn.execute("PRAGMA table_info(leads)").fetchall()
        }
        for column in EXPORT_COLUMNS:
            if column not in existing:
                conn.execute(f"ALTER TABLE leads ADD COLUMN {column} TEXT")
        conn.commit()

    def generate_filename(self, prefix: str, extension: str) -> str:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return f"{prefix}_{timestamp}.{extension}"

    def _prepare_record(self, record: LeadRecord) -> LeadRecord:
        prepared = replace(record).normalize()
        prepared.id = prepared.stable_id()
        return prepared

    def _serialize_config(self, job: ScrapingJob) -> dict:
        if not job.config:
            return {}
        data = dict(job.config.__dict__)
        source_type = data.get("source_type")
        if source_type is not None and hasattr(source_type, "value"):
            data["source_type"] = source_type.value
        return data

    def save_records_to_db(self, records: list[LeadRecord], db_path: str) -> int:
        """Persist imported or scraped records into SQLite leads table."""
        if not records:
            return 0
        conn = sqlite3.connect(db_path, timeout=30.0)
        try:
            conn.execute("PRAGMA journal_mode=WAL")
            conn.execute("PRAGMA synchronous=NORMAL")
            self._init_db(conn)
            count = 0
            for record in records:
                prepared = self._prepare_record(record)
                row = prepared.to_dict()
                conn.execute(
                    f"INSERT OR REPLACE INTO leads ({','.join(EXPORT_COLUMNS)}) VALUES ({','.join(['?']*len(EXPORT_COLUMNS))})",
                    [row.get(c, "") for c in EXPORT_COLUMNS],
                )
                count += 1
            conn.commit()
            return count
        finally:
            conn.close()

    def import_spreadsheet(self, path: str) -> list[LeadRecord]:
        """
        Parse an Excel (.xlsx, .xls) or CSV spreadsheet into LeadRecord instances,
        auto-mapping standard German/English headers (GmBh, ansprechner, email, etc.).
        """
        import uuid
        import re
        from datetime import datetime, timezone
        p = Path(path)
        if not p.exists():
            return []
        suffix = p.suffix.lower()
        rows: list[list[str]] = []

        if suffix == ".csv":
            with open(path, newline="", encoding="utf-8-sig", errors="replace") as f:
                reader = csv.reader(f)
                rows = [list(map(str, row)) for row in reader]
        else:
            import importlib.util
            if importlib.util.find_spec("openpyxl") is not None and suffix == ".xlsx":
                import openpyxl
                wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
                ws = wb.active
                for row in ws.iter_rows(values_only=True):
                    rows.append([str(c) if c is not None else "" for c in row])
            elif importlib.util.find_spec("xlrd") is not None:
                import xlrd
                wb = xlrd.open_workbook(path)
                ws = wb.sheet_by_index(0)
                for rx in range(ws.nrows):
                    rows.append([str(ws.cell(rx, cx).value) for cx in range(ws.ncols)])
            else:
                raise RuntimeError("Excel import requires openpyxl or xlrd.")

        if not rows:
            return []

        raw_headers = [str(col or "").strip() for col in rows[0]]
        norm_headers = [re.sub(r"[^a-z0-9]", "", h.lower()) for h in raw_headers]

        col_map: dict[str, int] = {}
        def match_col(field_name: str, keywords: list[str]):
            for idx, nh in enumerate(norm_headers):
                if idx in col_map.values():
                    continue
                if any(kw in nh for kw in keywords):
                    col_map[field_name] = idx
                    return

        match_col("company_name", ["company", "firma", "gmbh", "name", "unternehmen", "einrichtung", "klinik", "heim", "betreiber", "organisation", "employer", "arbeitgeber"])
        match_col("email", ["email", "e-mail", "mail", "e_mail", "kontakt_email", "emailadresse"])
        match_col("contact_person", ["ansprech", "contact", "person", "hr", "leiter", "recipient", "empfnger", "empfaenger", "anrede"])
        match_col("phone", ["phone", "telefon", "tel", "mobil", "rufnummer"])
        match_col("website", ["website", "url", "web", "homepage", "site"])
        match_col("city", ["city", "stadt", "ort", "location"])
        match_col("postal_code", ["plz", "postal", "zip"])
        match_col("address", ["address", "adresse", "str", "strasse", "strae"])
        match_col("job_title", ["job", "stelle", "titel", "title", "position", "beruf", "ausschreibung"])

        records: list[LeadRecord] = []
        now_ts = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S")

        for idx in range(1, len(rows)):
            row = rows[idx]
            if not any(str(c).strip() for c in row):
                continue
            def get_val(key: str) -> str:
                cidx = col_map.get(key)
                if cidx is not None and cidx < len(row):
                    val = str(row[cidx]).strip()
                    return val if val != "None" else ""
                return ""

            email = get_val("email").lower()
            company_name = get_val("company_name")
            contact_person = get_val("contact_person")
            phone = get_val("phone")
            website = get_val("website")
            city = get_val("city")
            postal_code = get_val("postal_code")
            address = get_val("address")
            job_title = get_val("job_title")

            if not email:
                for c in row:
                    s = str(c).strip()
                    if "@" in s and "." in s and len(s) < 100 and " " not in s:
                        email = s.lower()
                        break

            if not email and not company_name:
                continue

            if not company_name and email:
                dom = email.split("@")[-1].split(".")[0].capitalize()
                company_name = dom

            rec = LeadRecord(
                id="",
                source_type=SourceType.FILE,
                company_name=company_name,
                email=email,
                contact_person=contact_person,
                phone=phone,
                website=website,
                city=city,
                postal_code=postal_code,
                address=address,
                job_title=job_title or "Initiativbewerbung",
                scraped_at=now_ts,
            )
            rec.id = rec.stable_id()
            records.append(rec)

        return records

# 1.1.0 Beta5
